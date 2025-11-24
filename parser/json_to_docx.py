#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для заполнения Word документа данными из JSON файла.
Сохраняет форматирование и автоматически клонирует блоки для множественных записей.
"""

import sys
import os
import json
import argparse
import re
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Ошибка: библиотека python-docx не установлена.")
    print("Установите её командой: pip install python-docx")
    sys.exit(1)

from docx.oxml.ns import qn


# Карты месяцев и специальных слов для вычисления стажа и форматирования периодов
MONTHS_MAP = {
    'январь': 1, 'января': 1,
    'февраль': 2, 'февраля': 2,
    'март': 3, 'марта': 3,
    'апрель': 4, 'апреля': 4,
    'май': 5, 'мая': 5,
    'июнь': 6, 'июня': 6,
    'июль': 7, 'июля': 7,
    'август': 8, 'августа': 8,
    'сентябрь': 9, 'сентября': 9,
    'октябрь': 10, 'октября': 10,
    'ноябрь': 11, 'ноября': 11,
    'декабрь': 12, 'декабря': 12,
    'january': 1, 'february': 2, 'march': 3, 'april': 4,
    'may': 5, 'june': 6, 'july': 7, 'august': 8,
    'september': 9, 'october': 10, 'november': 11, 'december': 12
}

CURRENT_PERIOD_TERMS = ['настоящее время', 'по настоящее время', 'по наст. время', 'н.в.', 'present', 'current']

DURATION_WORD_REPLACEMENTS = {
    'год': 'ГОД',
    'года': 'ГОДА',
    'году': 'ГОДУ',
    'годом': 'ГОДОМ',
    'лет': 'ЛЕТ',
    'г.': 'Г.',
    'г': 'Г',
    'месяц': 'МЕСЯЦ',
    'месяца': 'МЕСЯЦА',
    'месяцев': 'МЕСЯЦЕВ',
    'месяце': 'МЕСЯЦЕ',
    'мес': 'МЕС',
    'мес.': 'МЕС.',
}

DURATION_WORD_PATTERN = re.compile(r'\b(' + '|'.join(re.escape(k) for k in DURATION_WORD_REPLACEMENTS.keys()) + r')\b', re.IGNORECASE)

DEFAULT_FONT_NAME = "Calibri Light"
DEFAULT_FONT_SIZE_PT = 10.5
BULLET_LEFT_INDENT_CM = 0.63


def apply_default_font(run):
    """Применяет шрифт Calibri Light 10.5 к run."""
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE_PT)
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)


def iter_container_paragraphs(container):
    """Итерирует все параграфы в документе/ячейке, включая вложенные таблицы."""
    if hasattr(container, 'paragraphs'):
        for paragraph in container.paragraphs:
            yield paragraph
    if hasattr(container, 'tables'):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_container_paragraphs(cell)


def apply_default_font_to_document(doc):
    """Проходит по всему документу и приводит текст к Calibri Light."""
    for paragraph in iter_container_paragraphs(doc):
        for run in paragraph.runs:
            apply_default_font(run)


def add_run_with_default_font(paragraph, text):
    run = paragraph.add_run(text)
    apply_default_font(run)
    return run


def configure_bullet_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.left_indent = Cm(BULLET_LEFT_INDENT_CM)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    remove_paragraph_numbering(paragraph)


def ensure_runs_not_bold(paragraph):
    for run in paragraph.runs:
        run.font.bold = False
        run.bold = False


def write_label_and_value(paragraph, label_text, value_text):
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    if label_text:
        if template_run:
            label_run = paragraph.add_run(label_text)
            clone_run_formatting(template_run, label_run)
        else:
            label_run = add_run_with_default_font(paragraph, label_text)
        if not label_text.endswith(' '):
            if template_run:
                spacer = paragraph.add_run(' ')
                clone_run_formatting(template_run, spacer)
            else:
                add_run_with_default_font(paragraph, ' ')
    if value_text:
        add_run_with_default_font(paragraph, value_text)


def normalize_label_value_format(paragraph, template_para=None):
    """Делит параграф на метку и значение, оставляя значение без жирного начертания."""
    full_text = paragraph.text or ""
    if ':' not in full_text:
        return
    colon_idx = full_text.find(':')
    label_text = full_text[:colon_idx + 1].rstrip()
    remainder = full_text[colon_idx + 1:]
    value_text = remainder.strip()

    template_run = None
    if template_para and template_para.runs:
        template_run = template_para.runs[0]

    paragraph.clear()
    if label_text:
        if template_run:
            label_run = paragraph.add_run(label_text)
            clone_run_formatting(template_run, label_run)
        else:
            add_run_with_default_font(paragraph, label_text)
        if value_text:
            add_run_with_default_font(paragraph, ' ')
    if value_text:
        add_run_with_default_font(paragraph, value_text)


def normalize_bullet_items(items, placeholders=None):
    placeholders = placeholders or []
    normalized = []
    for item in items or []:
        text = format_list_item(item) if isinstance(item, dict) else str(item)
        text = text.replace('•', '').strip()
        if text and text not in placeholders:
            normalized.append(text)
    return normalized


def set_bullet_list_in_cell(cell, items):
    if cell is None:
        return False
    normalized = normalize_bullet_items(items)
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    target_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
    target_para.clear()
    if not normalized:
        remove_paragraph_numbering(target_para)
        return False
    for idx, text in enumerate(normalized):
        current_para = target_para if idx == 0 else cell.add_paragraph("")
        configure_bullet_paragraph(current_para)
        add_run_with_default_font(current_para, f"• {text}")
    return True


def set_labeled_bullet_list(cell, fallback_label, items):
    if cell is None:
        return False
    normalized = normalize_bullet_items(items)
    if not cell.paragraphs:
        cell.add_paragraph("")
    label_para = cell.paragraphs[0]
    label_run = label_para.runs[0] if label_para.runs else None
    label_text = fallback_label
    para_text = label_para.text.strip()
    colon_idx = para_text.find(':')
    if colon_idx != -1:
        label_text = para_text[:colon_idx + 1]
    label_para.clear()
    if label_run:
        new_label = label_para.add_run(label_text)
        clone_run_formatting(label_run, new_label)
    else:
        add_run_with_default_font(label_para, label_text)
    # Удаляем остальные параграфы
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    if not normalized:
        return False
    for text in normalized:
        bullet_para = cell.add_paragraph("")
        configure_bullet_paragraph(bullet_para)
        add_run_with_default_font(bullet_para, f"• {text}")
    return True


def set_bullet_list_in_document(doc, indices, items):
    if not indices:
        return False
    normalized = normalize_bullet_items(items)
    if not normalized:
        return False
    sorted_indices = sorted(set(idx for idx in indices if idx is not None))
    if not sorted_indices:
        return False
    base_idx = sorted_indices[0]
    for idx in reversed(sorted_indices[1:]):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para._element.getparent().remove(para._element)
    if base_idx >= len(doc.paragraphs):
        return False
    current_para = doc.paragraphs[base_idx]
    current_para.clear()
    for i, text in enumerate(normalized):
        if i > 0:
            current_para = current_para.insert_paragraph_after()
        configure_bullet_paragraph(current_para)
        add_run_with_default_font(current_para, f"• {text}")
    return True


def load_json(json_path):
    """
    Загружает JSON файл.
    
    Args:
        json_path (str): Путь к JSON файлу
        
    Returns:
        dict: Данные из JSON
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Ошибка при чтении JSON файла: {e}")
        sys.exit(1)


def find_placeholder_runs(paragraph, placeholder):
    """
    Находит все runs в параграфе, содержащие плейсхолдер.
    
    Args:
        paragraph: Параграф документа
        placeholder (str): Текст плейсхолдера (например, "{{vacancy}}")
        
    Returns:
        list: Список индексов runs, содержащих плейсхолдер
    """
    indices = []
    text = ""
    for i, run in enumerate(paragraph.runs):
        text += run.text
        if placeholder in text:
            indices.append(i)
    return indices


def replace_text_preserving_format(paragraph, old_text, new_text, force_default_font=True):
    """
    Заменяет текст в параграфе, сохраняя форматирование.
    Использует простой подход: заменяет весь текст параграфа, сохраняя форматирование первого run.
    
    Args:
        paragraph: Параграф документа
        old_text (str): Текст для замены
        new_text (str): Новый текст
        
    Returns:
        bool: True если замена выполнена успешно
    """
    # Проверяем, есть ли текст в параграфе
    full_text = paragraph.text
    if old_text not in full_text:
        return False
    
    # Если текст пустой, просто удаляем плейсхолдер
    if not new_text:
        new_text = ""
    
    # Сохраняем форматирование первого run (если есть)
    font_name = None
    font_size = None
    font_bold = None
    font_italic = None
    font_underline = None
    font_color_rgb = None
    
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_underline = first_run.font.underline
        # Сохраняем цвет правильно (цвет может быть RGB или theme_color)
        if first_run.font.color and first_run.font.color.rgb:
            font_color_rgb = first_run.font.color.rgb
    
    # Заменяем текст в параграфе
    new_paragraph_text = full_text.replace(old_text, new_text)
    
    # Если в параграфе нет runs, создаем один
    if not paragraph.runs:
        paragraph.add_run("")
    
    # Используем первый run и сохраняем его форматирование
    target_run = paragraph.runs[0]
    target_run.text = new_paragraph_text
    
    # Удаляем остальные runs, чтобы избежать дублирования текста
    for run in reversed(paragraph.runs[1:]):
        paragraph._element.remove(run._element)
    
    # Применяем форматирование
    if font_name:
        target_run.font.name = font_name
    if font_size:
        target_run.font.size = font_size
    if font_bold is not None:
        target_run.font.bold = font_bold
    if font_italic is not None:
        target_run.font.italic = font_italic
    if font_underline is not None:
        target_run.font.underline = font_underline
    if font_color_rgb:
        target_run.font.color.rgb = font_color_rgb
    if force_default_font and new_text:
        apply_default_font(target_run)

    return True


def clone_paragraph_formatting(source_para, target_para):
    """
    Клонирует форматирование параграфа.
    
    Args:
        source_para: Исходный параграф
        target_para: Целевой параграф
    """
    target_para.style = source_para.style
    target_para.alignment = source_para.alignment
    target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing


def clone_run_formatting(source_run, target_run):
    """
    Клонирует форматирование run.
    
    Args:
        source_run: Исходный run
        target_run: Целевой run
    """
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic
    if source_run.font.underline is not None:
        target_run.font.underline = source_run.font.underline
    # Копируем цвет правильно
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def set_paragraph_text(paragraph, text, template_para=None):
    """
    Заменяет текст параграфа, очищая существующие run и при необходимости применяя форматирование.
    
    Args:
        paragraph: Параграф для обновления
        text (str): Новый текст
        template_para: Параграф-шаблон для копирования форматирования run (опционально)
    """
    for run in reversed(paragraph.runs):
        paragraph._element.remove(run._element)
    
    new_run = add_run_with_default_font(paragraph, text)
    if template_para and template_para.runs:
        clone_run_formatting(template_para.runs[0], new_run)


def uppercase_duration_words(text):
    """Выделяет слова про длительности (год, месяц) капсом."""
    if not text:
        return text

    def replacer(match):
        word = match.group(0)
        return DURATION_WORD_REPLACEMENTS.get(word.lower(), word.upper())

    return DURATION_WORD_PATTERN.sub(replacer, text)


def _parse_single_date(text_value):
    """Преобразует строку вида 'Январь 2020' или 'настоящее время' в date."""
    if not text_value:
        return None
    value = text_value.strip().lower()
    if not value:
        return None
    if any(term in value for term in CURRENT_PERIOD_TERMS):
        today = date.today()
        return date(today.year, today.month, 1)

    year_match = re.search(r'(19|20)\d{2}', value)
    if not year_match:
        return None
    year = int(year_match.group(0))
    month = 1
    for name, number in MONTHS_MAP.items():
        if name in value:
            month = number
            break
    return date(year, month, 1)


def parse_period_range(period_str):
    """Возвращает (start_date, end_date) из строки периода."""
    if not period_str:
        return (None, None)
    parts = re.split(r'[\u2013\u2014\-]+', period_str)
    start_part = parts[0].strip() if parts else period_str.strip()
    end_part = parts[1].strip() if len(parts) > 1 else ''
    start_date = _parse_single_date(start_part)
    if end_part:
        end_date = _parse_single_date(end_part)
        if not end_date:
            end_date = date.today()
    else:
        end_date = date.today()
    return (start_date, end_date)


def calculate_experience_months(work_experience):
    """Суммирует продолжительность всех мест работы в месяцах."""
    total_months = 0
    for item in work_experience or []:
        period = item.get('period', '')
        start_date, end_date = parse_period_range(period)
        if not start_date:
            continue
        if not end_date:
            end_date = date.today()
        months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
        if months < 0:
            continue
        total_months += max(months, 0)
    return total_months


def format_experience_summary(work_experience):
    """Возвращает строку с опытом вида '8 ЛЕТ 3 МЕСЯЦА'."""
    total_months = calculate_experience_months(work_experience)
    if total_months <= 0:
        return "МЕНЕЕ 1 МЕСЯЦА"
    years = total_months // 12
    months = total_months % 12
    parts = []
    if years:
        if years % 10 == 1 and years % 100 != 11:
            word = 'ГОД'
        elif 2 <= years % 10 <= 4 and not 12 <= years % 100 <= 14:
            word = 'ГОДА'
        else:
            word = 'ЛЕТ'
        parts.append(f"{years} {word}")
    if months:
        if months % 10 == 1 and months % 100 != 11:
            month_word = 'МЕСЯЦ'
        elif 2 <= months % 10 <= 4 and not 12 <= months % 100 <= 14:
            month_word = 'МЕСЯЦА'
        else:
            month_word = 'МЕСЯЦЕВ'
        parts.append(f"{months} {month_word}")
    return ' '.join(parts) if parts else "МЕНЕЕ 1 МЕСЯЦА"


def fill_label_paragraph(doc, label_variants, value, uppercase_value=False):
    """Находит параграф с меткой и подставляет значение после неё без наследования форматирования."""
    if not value:
        return False
    value_text = str(value).strip()
    if not value_text:
        return False
    if uppercase_value:
        value_text = value_text.upper()

    labels = label_variants if isinstance(label_variants, (list, tuple)) else [label_variants]
    labels_upper = [lbl.upper() for lbl in labels]

    def iter_paragraphs():
        for paragraph in doc.paragraphs:
            yield paragraph
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    for para in iter_paragraphs():
        para_text = para.text.strip()
        if not para_text:
            continue
        para_upper = para_text.upper()
        if any(para_upper.startswith(lbl) for lbl in labels_upper):
            colon_idx = para.text.find(':')
            if colon_idx != -1:
                label_text = para.text[:colon_idx + 1]
            else:
                label_text = para.text.strip()
            write_label_and_value(para, label_text, value_text)
            return True
    return False


def normalize_category_name(name):
    """Приводит название категории к формату: первое слово с заглавной буквы, остальные строчные."""
    if not name:
        return ''
    words = name.strip().split()
    if not words:
        return ''
    normalized = [words[0].capitalize()]
    normalized.extend(word.lower() for word in words[1:])
    return ' '.join(normalized)


def remove_paragraph_numbering(paragraph):
    """Убирает нумерацию/маркеры из параграфа."""
    pPr = paragraph._element.pPr
    if pPr is not None and pPr.numPr is not None:
        pPr.remove(pPr.numPr)


def find_value_cell_for_header(doc, header_keywords):
    """Возвращает ячейку таблицы с данными после заголовка."""
    keywords = header_keywords if isinstance(header_keywords, (list, tuple)) else [header_keywords]
    keywords = [kw.lower() for kw in keywords]
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            header_cell = row.cells[0]
            header_text = ' '.join(p.text.strip() for p in header_cell.paragraphs).strip().lower()
            if any(keyword in header_text for keyword in keywords):
                if len(row.cells) > 1:
                    return row.cells[1]
                return row.cells[0]
    return None


def fill_skills_section(doc, skills):
    """Заполняет блок Навыки и инструменты без маркеров, делая категории жирными."""
    if not skills:
        return False
    cell = find_value_cell_for_header(doc, ['навыки и инструменты', 'skills and tools'])
    if cell is None:
        return False

    # Готовим ячейку: оставляем только первый параграф
    first_para = cell.paragraphs[0]
    first_para.clear()
    remove_paragraph_numbering(first_para)
    # Удаляем остальные параграфы
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)

    added = 0
    for idx, item in enumerate(skills):
        if isinstance(item, dict):
            item_text = format_list_item(item)
        else:
            item_text = str(item)
        item_text = item_text.strip()
        if not item_text:
            continue

        if ':' in item_text:
            category, details = item_text.split(':', 1)
        else:
            category, details = item_text, ''
        category = normalize_category_name(category.strip())
        details = details.strip()

        para = first_para if added == 0 else cell.add_paragraph()
        remove_paragraph_numbering(para)
        if category:
            label_text = category + (':' if details else '')
            category_run = add_run_with_default_font(para, label_text)
            category_run.bold = True
            if details:
                add_run_with_default_font(para, f" {details}")
        else:
            add_run_with_default_font(para, details)
        added += 1

    return added > 0


def format_list_item(item):
    """
    Форматирует элемент списка в строку.
    Если элемент - словарь (например, образование), форматирует его в читаемый текст.
    
    Args:
        item: Элемент списка (строка или словарь)
        
    Returns:
        str: Отформатированная строка
    """
    if isinstance(item, dict):
        # Обработка словарей (например, образование)
        if 'degree' in item or 'institution' in item:
            # Формат для образования
            parts = []
            if item.get('degree'):
                parts.append(item['degree'])
            if item.get('specialization'):
                parts.append(f"специальность: {item['specialization']}")
            if item.get('institution'):
                parts.append(item['institution'])
            if item.get('faculty'):
                parts.append(item['faculty'])
            if item.get('year'):
                parts.append(f"({item['year']})")
            return ", ".join(parts) if parts else str(item)
        # Для других словарей просто объединяем значения
        return ", ".join([str(v) for v in item.values() if v]) if item else str(item)
    return str(item)


def find_template_block(doc, start_marker, end_marker):
    """Находит блок параграфов между маркерами."""
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        if start_marker in para.text:
            start_idx = i
        if end_idx is None and start_idx is not None and end_marker in para.text:
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        return (start_idx, end_idx, doc.paragraphs[start_idx:end_idx + 1])
    return None


def process_simple_fields(doc, data):
    """Обрабатывает простые поля (не массивы)."""
    simple_fields = {
        'vacancy': data.get('vacancy', ''),
        'pitch': data.get('pitch', ''),
    }
    general_info = data.get('general_info', {})
    simple_fields.update({
        'foreign_language': general_info.get('foreign_language', ''),
        'citizenship_location': general_info.get('citizenship_location', ''),
        'employment': general_info.get('employment', ''),
        'status': general_info.get('status', ''),
    })

    replaced_count = 0
    for field_name, field_value in simple_fields.items():
        placeholder = f"{{{{{field_name}}}}}"
        value_str = str(field_value) if field_value else ""

        for para in doc.paragraphs:
            if placeholder in para.text and replace_text_preserving_format(para, placeholder, value_str):
                replaced_count += 1
                print(f"  ✓ {field_name}: {value_str[:50] if value_str else '(пусто)'}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text and replace_text_preserving_format(para, placeholder, value_str):
                            replaced_count += 1
                            print(f"  ✓ {field_name} (в таблице): {value_str[:50] if value_str else '(пусто)'}")

    return replaced_count


def process_list_field(doc, data, field_path, placeholder_name):
    """Обрабатывает поле-список (массив строк)."""
    value = data
    for key in field_path:
        value = value.get(key, {})
    if not isinstance(value, list):
        return 0

    start_marker = f"{{{{#{placeholder_name}}}}}"
    end_marker = f"{{{{/{placeholder_name}}}}}"
    block = find_template_block(doc, start_marker, end_marker)
    if not block:
        placeholder = f"{{{{{placeholder_name}}}}}"
        found = False
        for para in doc.paragraphs:
            if placeholder in para.text:
                if placeholder_name == 'skills_and_tools' and value and any(':' in str(item) for item in value):
                    list_text = "\n\n".join([format_list_item(item) for item in value]) if value else ""
                else:
                    list_text = "\n".join([f"• {format_list_item(item)}" for item in value]) if value else ""
                if replace_text_preserving_format(para, placeholder, list_text):
                    found = True
        if found:
            print(f"  ✓ {placeholder_name}: {len(value)} элементов (простой плейсхолдер)")
            return 1
        return 0

    start_idx, end_idx, template_paras = block
    for para in template_paras:
        if start_marker in para.text:
            replace_text_preserving_format(para, start_marker, "")
        if end_marker in para.text:
            replace_text_preserving_format(para, end_marker, "")

    if value and len(value) > 0:
        template_para = None
        for para in template_paras:
            if start_marker not in para.text and end_marker not in para.text and para.text.strip():
                template_para = para
                break

        if template_para:
            for i in range(end_idx - 1, start_idx, -1):
                if i < len(doc.paragraphs):
                    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

            insert_idx = start_idx + 1
            is_with_categories = placeholder_name == 'skills_and_tools' and value and any(':' in str(item) for item in value)

            for item in value:
                new_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                clone_paragraph_formatting(template_para, new_para)
                formatted_item = format_list_item(item)

                if is_with_categories and ':' in str(formatted_item):
                    if template_para.runs:
                        for source_run in template_para.runs:
                            new_run = new_para.add_run(formatted_item)
                            clone_run_formatting(source_run, new_run)
                    else:
                        new_para.add_run(formatted_item)
                else:
                    if template_para.runs:
                        for source_run in template_para.runs:
                            template_text = template_para.text.strip()
                            text_value = f"• {formatted_item}" if template_text.startswith(('•', '-')) else formatted_item
                            new_run = new_para.add_run(text_value)
                            clone_run_formatting(source_run, new_run)
                    else:
                        new_para.add_run(formatted_item)

                insert_idx += 1
            print(f"  ✓ {placeholder_name}: {len(value)} элементов")
            return 1
    return 0


def process_work_experience(doc, data):
    """Обрабатывает блок опыта работы."""
    work_experience = data.get('work_experience', [])
    if not work_experience:
        return 0

    start_marker = "{{#work_experience}}"
    end_marker = "{{/work_experience}}"
    block = find_template_block(doc, start_marker, end_marker)
    if not block:
        print(f"  ⚠️  Блок {start_marker}...{end_marker} не найден в шаблоне")
        return 0

    start_idx, end_idx, template_paras = block
    template_paras_clean = [p for p in template_paras if start_marker not in p.text and end_marker not in p.text]
    if not template_paras_clean:
        return 0

    for para in doc.paragraphs:
        if start_marker in para.text:
            replace_text_preserving_format(para, start_marker, "")
        if end_marker in para.text:
            replace_text_preserving_format(para, end_marker, "")

    for i in range(end_idx - 1, start_idx, -1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    insert_idx = start_idx + 1
    added_count = 0
    for work_item in work_experience:
        for template_para in template_paras_clean:
            new_para = doc.paragraphs[insert_idx].insert_paragraph_before()
            clone_paragraph_formatting(template_para, new_para)

            para_text = template_para.text
            replacements = {
                '{{company}}': work_item.get('company', ''),
                '{{position}}': work_item.get('position', ''),
                '{{period}}': work_item.get('period', ''),
            }
            for placeholder, value in replacements.items():
                if placeholder in para_text:
                    para_text = para_text.replace(placeholder, str(value))

            set_paragraph_text(new_para, para_text, template_para)

            if '{{responsibilities}}' in para_text:
                responsibilities = work_item.get('responsibilities', [])
                list_text = "\n".join([f"• {item}" for item in responsibilities]) if responsibilities else ""
                replace_text_preserving_format(new_para, '{{responsibilities}}', list_text)

            if '{{technologies}}' in para_text:
                technologies = work_item.get('technologies', [])
                if technologies:
                    tech_text = '\n'.join(technologies) if any(':' in t for t in technologies) else ", ".join(technologies)
                else:
                    tech_text = ""
                replace_text_preserving_format(new_para, '{{technologies}}', tech_text)

            insert_idx += 1
        added_count += 1
        print(f"  ✓ Опыт работы: {work_item.get('company', 'Не указано')} - {work_item.get('position', 'Не указано')}")

    return added_count


def process_project_experience(doc, data):
    """Обрабатывает блок проектного опыта."""
    project_experience = data.get('project_experience', [])
    if not project_experience:
        return 0

    start_marker = "{{#project_experience}}"
    end_marker = "{{/project_experience}}"
    block = find_template_block(doc, start_marker, end_marker)
    if not block:
        print(f"  ⚠️  Блок {start_marker}...{end_marker} не найден в шаблоне")
        return 0

    start_idx, end_idx, template_paras = block
    template_paras_clean = [p for p in template_paras if start_marker not in p.text and end_marker not in p.text]
    if not template_paras_clean:
        return 0

    for para in doc.paragraphs:
        if start_marker in para.text:
            replace_text_preserving_format(para, start_marker, "")
        if end_marker in para.text:
            replace_text_preserving_format(para, end_marker, "")

    for i in range(end_idx - 1, start_idx, -1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    insert_idx = start_idx + 1
    added_count = 0
    for project_item in project_experience:
        for template_para in template_paras_clean:
            new_para = doc.paragraphs[insert_idx].insert_paragraph_before()
            clone_paragraph_formatting(template_para, new_para)

            para_text = template_para.text
            template_text = template_para.text
            has_role_placeholder = "{{role}}" in template_text
            has_tech_placeholder = "{{technologies_and_tools}}" in template_text
            replacements = {
                '{{company}}': project_item.get('company', ''),
                '{{role}}': project_item.get('role', ''),
            }
            for placeholder, value in replacements.items():
                if placeholder in para_text:
                    para_text = para_text.replace(placeholder, str(value))

            set_paragraph_text(new_para, para_text, template_para)
            if has_role_placeholder:
                normalize_label_value_format(new_para, template_para)

            if '{{tasks}}' in para_text:
                tasks = project_item.get('tasks', [])
                tasks_text = "\n".join([f"• {item}" for item in tasks]) if tasks else ""
                replace_text_preserving_format(new_para, '{{tasks}}', tasks_text)

            if '{{technologies_and_tools}}' in para_text:
                tech = project_item.get('technologies_and_tools', [])
                if tech:
                    tech_text = '\n'.join(tech) if any(':' in t for t in tech) else ", ".join(tech)
                else:
                    tech_text = ""
                replace_text_preserving_format(new_para, '{{technologies_and_tools}}', tech_text)
                if has_tech_placeholder:
                    normalize_label_value_format(new_para, template_para)

            insert_idx += 1
        added_count += 1
        print(f"  ✓ Проект: {project_item.get('company', 'Не указано')} - {project_item.get('role', 'Не указано')}")

    return added_count


def find_section_by_header(doc, header_keywords, search_in_tables=True):
    """
    Находит секцию документа по заголовку.
    
    Args:
        doc: Документ
        header_keywords (list): Список ключевых слов для поиска заголовка
        search_in_tables (bool): Искать ли в таблицах
        
    Returns:
        tuple: (тип, индекс) где тип - 'paragraph' или 'table', индекс - позиция, или None
    """
    # Ищем в параграфах
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower().strip()
        for keyword in header_keywords:
            if keyword.lower() in text_lower:
                # Возвращаем следующий параграф после заголовка
                return ('paragraph', i + 1 if i + 1 < len(doc.paragraphs) else i)
    
    # Ищем в таблицах
    if search_in_tables:
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        text_lower = para.text.lower().strip()
                        for keyword in header_keywords:
                            if keyword.lower() in text_lower:
                                return ('table', (table_idx, row_idx, cell_idx))
    
    return None


def find_empty_paragraph_after_header(doc, header_keywords, max_search=15):
    """
    Находит пустой параграф или параграф с плейсхолдером после заголовка.
    
    Args:
        doc: Документ
        header_keywords (list): Ключевые слова заголовка
        max_search (int): Максимальное количество параграфов для поиска
        
    Returns:
        tuple: (тип, объект) где тип - 'paragraph' или 'table_cell', объект - параграф или ячейка таблицы
    """
    header_info = find_section_by_header(doc, header_keywords)
    if header_info is None:
        return None
    
    header_type, header_idx = header_info
    
    if header_type == 'paragraph':
        # Ищем пустой параграф или параграф с плейсхолдером, пропуская заголовки
        for i in range(header_idx, min(header_idx + max_search, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Пропускаем заголовки
            text_lower = text.lower()
            is_header = any(kw.lower() in text_lower for kw in header_keywords)
            if is_header and text:  # Если это заголовок и он не пустой, пропускаем
                continue
            
            # Если параграф пустой или содержит только пробелы/спецсимволы
            if not text or text in ['', '—', '-', '•', 'Место для указания вакансии']:
                return ('paragraph', para)
            # Если содержит плейсхолдер
            if '{{' in text:
                return ('paragraph', para)
            # Если это не заголовок и не пустой, используем его (но проверяем, не похож ли он на заголовок других секций)
            if text and not is_header:
                # Проверяем, не является ли это заголовком другой секции
                other_section_keywords = ['опыт работы', 'проектный опыт', 'общая информация', 'скрининг', 
                                         'образование', 'навыки', 'вакансия', 'work experience', 
                                         'project experience', 'general info', 'screening']
                is_other_header = any(kw.lower() in text_lower for kw in other_section_keywords)
                if not is_other_header:
                    return ('paragraph', para)
        
        # Если не нашли, возвращаем первый параграф после заголовка (если он не заголовок)
        if header_idx < len(doc.paragraphs):
            para = doc.paragraphs[header_idx]
            text_lower = para.text.strip().lower()
            is_header = any(kw.lower() in text_lower for kw in header_keywords)
            if not is_header:
                return ('paragraph', para)
    
    elif header_type == 'table':
        table_idx, row_idx, cell_idx = header_idx
        table = doc.tables[table_idx]
        header_cell = table.rows[row_idx].cells[cell_idx]
        
        # Проверяем, не является ли найденная ячейка заголовком
        header_text = ' '.join([p.text.strip() for p in header_cell.paragraphs]).lower()
        is_header = any(kw.lower() in header_text for kw in header_keywords)
        row = table.rows[row_idx]
        single_cell_row = len(row.cells) == 1
        if is_header and single_cell_row and header_cell.paragraphs:
            return ('table_cell', header_cell.paragraphs[0])

        if is_header and not single_cell_row:
            # Если это заголовок, ищем следующую ячейку в строке или следующую строку
            # Сначала пробуем следующую ячейку в той же строке
            if cell_idx + 1 < len(table.rows[row_idx].cells):
                next_cell = table.rows[row_idx].cells[cell_idx + 1]
                for para in next_cell.paragraphs:
                    text = para.text.strip()
                    if not text or text in ['', '—', '-', '•'] or '{{' in text:
                        return ('table_cell', para)
                # Если ячейка не содержит пустых параграфов, используем первый или создаем новый
                if next_cell.paragraphs:
                    return ('table_cell', next_cell.paragraphs[0])
                else:
                    new_para = next_cell.add_paragraph("")
                    return ('table_cell', new_para)
            
            # Если следующей ячейки нет, ищем следующую строку
            if row_idx + 1 < len(table.rows):
                next_row_cell = table.rows[row_idx + 1].cells[0] if table.rows[row_idx + 1].cells else None
                if next_row_cell:
                    for para in next_row_cell.paragraphs:
                        text = para.text.strip()
                        if not text or text in ['', '—', '-', '•'] or '{{' in text:
                            return ('table_cell', para)
                    if next_row_cell.paragraphs:
                        return ('table_cell', next_row_cell.paragraphs[0])
                    else:
                        new_para = next_row_cell.add_paragraph("")
                        return ('table_cell', new_para)
        
        # Ищем в самой ячейке (если это не заголовок)
        for para in header_cell.paragraphs:
            text = para.text.strip()
            # Пропускаем заголовки
            if any(kw.lower() in text.lower() for kw in header_keywords):
                continue
            if not text or text in ['', '—', '-', '•'] or '{{' in text:
                return ('table_cell', para)
        
        # Если не нашли, создаем новый параграф в той же ячейке после заголовка
        return ('table_cell', header_cell.add_paragraph(""))
    
    return None


def fill_by_header(doc, header_keywords, value, field_name, debug=False):
    """
    Заполняет поле, ища его по заголовку.
    
    Args:
        doc: Документ
        header_keywords (list): Ключевые слова заголовка
        value: Значение для заполнения
        field_name (str): Имя поля (для отладки)
        debug (bool): Выводить отладочную информацию
        
    Returns:
        bool: True если заполнение выполнено
    """
    target_info = find_empty_paragraph_after_header(doc, header_keywords)
    if target_info is None:
        if debug:
            print(f"  ⚠️  Не найден заголовок для поля '{field_name}' (ключевые слова: {header_keywords})")
        return False
    
    target_type, para = target_info
    value_str = str(value) if value else ""
    
    if debug:
        print(f"  🔍 Найден параграф для '{field_name}': '{para.text[:50]}' (тип: {target_type})")
    placeholder_texts = ['', '—', '-', '•', 'Место для указания вакансии', 'Рассказ о себе от первого лица']
    
    # Если параграф содержит плейсхолдер, заменяем его
    if '{{' in para.text:
        placeholder = re.search(r'\{\{[^}]+\}\}', para.text)
        if placeholder:
            if debug:
                print(f"  🔄 Замена плейсхолдера '{placeholder.group()}' на '{value_str[:50]}'")
            return replace_text_preserving_format(para, placeholder.group(), value_str)
    
    # Иначе заменяем весь текст параграфа
    old_text = para.text.strip()
    
    # Проверяем, не является ли найденный текст заголовком
    old_text_lower = old_text.lower()
    is_header = any(kw.lower() in old_text_lower for kw in header_keywords)
    
    if is_header:
        if target_type == 'table_cell':
            cell = para._parent
            table = getattr(cell, '_parent', None)
            row = None
            row_idx = None
            col_idx = 0
            if table is not None:
                for idx, tbl_row in enumerate(table.rows):
                    for c_idx, candidate_cell in enumerate(tbl_row.cells):
                        if candidate_cell._tc is cell._tc:
                            row = tbl_row
                            row_idx = idx
                            col_idx = c_idx
                            break
                    if row is not None:
                        break
            cells_in_row = len(row.cells) if row is not None else 1
            if cells_in_row == 1:
                target_para = None
                if table is not None and row_idx is not None and row_idx + 1 < len(table.rows):
                    next_row = table.rows[row_idx + 1]
                    if col_idx < len(next_row.cells):
                        next_cell = next_row.cells[col_idx]
                        next_text = ' '.join([p.text.strip() for p in next_cell.paragraphs]).strip()
                        if not next_text or next_text in placeholder_texts or '{{' in next_text:
                            if next_cell.paragraphs:
                                target_para = next_cell.paragraphs[0]
                            else:
                                target_para = next_cell.add_paragraph("")
                if target_para:
                    target_para.clear()
                    add_run_with_default_font(target_para, value_str)
                    if debug:
                        print(f"  ✅ Добавлен текст под заголовком в следующей строке: '{value_str[:50]}'")
                    return True
                para.clear()
                add_run_with_default_font(para, value_str)
                if debug:
                    print(f"  ✅ Заголовок заменен значением в той же ячейке: '{value_str[:50]}'")
                return True
            # Пытаемся найти дополнительный параграф в этой же ячейке
            target_para = None
            for idx, cell_para in enumerate(cell.paragraphs):
                if cell_para == para:
                    continue
                text = cell_para.text.strip()
                if not text or text in placeholder_texts or '{{' in text:
                    target_para = cell_para
                    break
            if target_para is None:
                target_para = cell.add_paragraph("")
            target_para.clear()
            add_run_with_default_font(target_para, value_str)
            if debug:
                print(f"  ✅ Добавлен текст в новую строку ячейки: '{value_str[:50]}'")
            return True
        # Если это заголовок в параграфе, ищем следующий параграф
        if debug:
            print(f"  ⚠️  Найденный текст является заголовком, ищем следующий параграф: '{old_text[:50]}'")
        # Пытаемся найти следующий параграф после заголовка
        para_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p == para:
                para_idx = i
                break
        
        if para_idx is not None and para_idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[para_idx + 1]
            next_text = next_para.text.strip()
            next_text_lower = next_text.lower()
            if not any(kw.lower() in next_text_lower for kw in header_keywords):
                other_section_keywords = ['опыт работы', 'проектный опыт', 'общая информация', 'скрининг', 
                                         'образование', 'навыки', 'вакансия', 'work experience', 
                                         'project experience', 'general info', 'screening']
                is_other_header = any(kw.lower() in next_text_lower for kw in other_section_keywords)
                if not is_other_header:
                    if next_text in placeholder_texts or '{{' in next_text:
                        next_para.clear()
                        add_run_with_default_font(next_para, value_str)
                        if debug:
                            print(f"  ✅ Заполнен следующий параграф: '{value_str[:50]}'")
                        return True
                    else:
                        replace_text_preserving_format(next_para, next_text, value_str)
                        if debug:
                            print(f"  ✅ Заменен текст в следующем параграфе: '{value_str[:50]}'")
                        return True
        
        return False
    
    # Специальная обработка для стандартных плейсхолдеров
    if old_text in ['Место для указания вакансии', '—', '-', '']:
        # Очищаем параграф и добавляем новый текст
        para.clear()
        add_run_with_default_font(para, value_str)
        if debug:
            print(f"  ✅ Заполнен пустой параграф: '{value_str[:50]}'")
        return True
    
    if old_text:
        if debug:
            print(f"  🔄 Замена текста '{old_text[:50]}' на '{value_str[:50]}'")
        result = replace_text_preserving_format(para, old_text, value_str)
        if not result:
            # Если замена не удалась, просто очищаем и добавляем новый текст
            para.clear()
            add_run_with_default_font(para, value_str)
            if debug:
                print(f"  ✅ Заменено через очистку: '{value_str[:50]}'")
            return True
        return result
    else:
        # Если параграф пустой, добавляем текст
        if para.runs:
            para.runs[0].text = value_str
            apply_default_font(para.runs[0])
        else:
            add_run_with_default_font(para, value_str)
        if debug:
            print(f"  ✅ Добавлен текст в пустой параграф: '{value_str[:50]}'")
        return True


def fill_document(template_path, json_data, output_path):
    """
    Заполняет документ данными из JSON.
    
    Args:
        template_path (str): Путь к шаблону Word
        json_data (dict): Данные из JSON
        output_path (str): Путь к выходному файлу
    """
    print(f"Открытие шаблона: {template_path}")
    doc = Document(template_path)
    
    # Подсчитываем найденные плейсхолдеры для отладки
    found_placeholders = []
    for para in doc.paragraphs:
        text = para.text
        if '{{' in text and '}}' in text:
            placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
            found_placeholders.extend(placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if '{{' in text and '}}' in text:
                        placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
                        found_placeholders.extend(placeholders)

    has_placeholders = len(found_placeholders) > 0

    if has_placeholders:
        print(f"Найдено плейсхолдеров в шаблоне: {len(set(found_placeholders))}")
        print(f"Уникальные плейсхолдеры: {', '.join(set(found_placeholders))}")
        print("\nИспользуется режим заполнения по плейсхолдерам...")
    else:
        print("⚠️  В шаблоне не найдено плейсхолдеров!")
        print("Переключение на режим заполнения по заголовкам...")
        print("\n📋 Структура документа для отладки:")
        print("-" * 60)
        for i, para in enumerate(doc.paragraphs[:20]):
            text = para.text.strip()
            if text:
                print(f"[{i:2d}] {text[:70]}")
        if len(doc.paragraphs) > 20:
            print(f"... и еще {len(doc.paragraphs) - 20} параграфов")
        print("-" * 60)

    if has_placeholders:
        print("\nЗаполнение простых полей...")
        replaced_simple = process_simple_fields(doc, json_data)
        print(f"  Заполнено простых полей: {replaced_simple}")

        print("\nЗаполнение списков...")
        replaced_lists = 0
        replaced_lists += process_list_field(doc, json_data, ['general_info', 'skills_and_tools'], 'skills_and_tools')
        replaced_lists += process_list_field(doc, json_data, ['general_info', 'education'], 'education')
        replaced_lists += process_list_field(doc, json_data, ['screening', 'hard_skills'], 'hard_skills')
        replaced_lists += process_list_field(doc, json_data, ['screening', 'soft_skills'], 'soft_skills')
        print(f"  Заполнено списков: {replaced_lists}")

        print("\nЗаполнение опыта работы...")
        work_count = process_work_experience(doc, json_data)
        print(f"  Добавлено записей опыта работы: {work_count}")

        print("\nЗаполнение проектного опыта...")
        project_count = process_project_experience(doc, json_data)
        print(f"  Добавлено записей проектного опыта: {project_count}")
    else:
        print("\nЗаполнение по заголовкам...")
        fill_by_headers_mode(doc, json_data, debug=True)

    apply_default_font_to_document(doc)
    
    print(f"\nСохранение документа: {output_path}")
    
    # Пробуем сохранить файл
    try:
        doc.save(output_path)
        print("✅ Документ успешно заполнен!")
    except PermissionError:
        # Если файл открыт, пробуем сохранить под другим именем
        import os
        from datetime import datetime
        
        base_name = os.path.splitext(output_path)[0]
        extension = os.path.splitext(output_path)[1]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_output_path = f"{base_name}_{timestamp}{extension}"
        
        print(f"⚠️  Файл {output_path} заблокирован (возможно, открыт в Word)")
        print(f"💾 Сохранение под новым именем: {new_output_path}")
        
        try:
            doc.save(new_output_path)
            print("✅ Документ успешно сохранен!")
            print(f"📄 Файл сохранен как: {new_output_path}")
        except Exception as e:
            print(f"❌ Ошибка при сохранении файла: {e}")
            print("💡 Закройте файл document_filled.docx и попробуйте снова")
            raise
    except Exception as e:
        print(f"❌ Ошибка при сохранении файла: {e}")
        raise


def fill_by_headers_mode(doc, json_data, debug=False):
    """
    Заполняет документ, ища данные по заголовкам.
    
    Args:
        doc: Документ
        json_data (dict): Данные из JSON
        debug (bool): Выводить отладочную информацию
    """
    replaced_count = 0
    work_experience = json_data.get('work_experience', [])
    general_info = json_data.get('general_info', {})
    
    # ФИО - простая замена текста "ФИО" на значение из JSON
    if json_data.get('full_name') or json_data.get('name'):
        full_name = json_data.get('full_name') or json_data.get('name', '')
        if full_name:
            # Ищем и заменяем текст "ФИО" во всех параграфах
            full_name_value = full_name.upper()
            for para in doc.paragraphs:
                if 'фио' in para.text.lower():
                    # Заменяем "ФИО" на значение
                    replace_text_preserving_format(para, 'ФИО', full_name_value)
                    replaced_count += 1
                    print(f"  ✓ ФИО: {full_name_value[:50]}")
                    break

    # Вакансия - простая замена через плейсхолдер или текст
    if json_data.get('vacancy'):
        vacancy_value = json_data['vacancy'].upper()
        if fill_by_header(doc, ['вакансия', 'vacancy', 'позиция', 'место для указания'], 
                         vacancy_value, 'vacancy', debug=debug):
            replaced_count += 1
            print(f"  ✓ Вакансия: {vacancy_value[:50]}")

    # ОПЫТ РАБОТЫ (суммарный стаж в годах/месяцах)
    experience_summary = format_experience_summary(work_experience)
    if experience_summary:
        if fill_label_paragraph(doc, 'ОПЫТ РАБОТЫ', experience_summary, uppercase_value=True):
            replaced_count += 1
            print(f"  ✓ ОПЫТ РАБОТЫ: {experience_summary}")

    # Проектный бекграунд (краткое описание)
    project_background = (
        json_data.get('project_background')
        or json_data.get('project_background_summary')
        or general_info.get('project_background')
    )
    if project_background:
        if fill_label_paragraph(doc, 'ПРОЕКТНЫЙ БЭКГРАУНД', project_background, uppercase_value=False):
            replaced_count += 1
            print(f"  ✓ Проектный бекграунд")
    
    # Питч
    if json_data.get('pitch'):
        # Пробуем разные варианты ключевых слов
        pitch_keywords_list = [
            ['питч', 'pitch'],
            ['рассказ о себе', 'рассказ', 'о себе'],
            ['питч:', 'pitch:'],
        ]
        found = False
        for pitch_keywords in pitch_keywords_list:
            if fill_by_header(doc, pitch_keywords, json_data['pitch'], 'pitch', debug=debug):
                replaced_count += 1
                print(f"  ✓ Питч: {json_data['pitch'][:50]}")
                found = True
                break
        if not found and debug:
            print(f"  ⚠️  Не найден заголовок для поля 'pitch'")
    
    # Общая информация
    
    # Навыки и инструменты
    skills = general_info.get('skills_and_tools', [])
    if skills:
        filled_skills = fill_skills_section(doc, skills)
        if not filled_skills:
            filled_skills = fill_list_by_header(
                doc,
                ['навыки и инструменты', 'skills and tools', 'skills_and_tools'],
                skills,
                'skills_and_tools',
                use_bullets=False
            )
        if filled_skills:
            replaced_count += 1
            print(f"  ✓ Навыки и инструменты: {len(skills)} элементов")

    # Образование
    education = general_info.get('education', [])
    if education:
        if fill_list_by_header(doc, ['образование', 'education'], education, 'education', use_bullets=False):
            replaced_count += 1
            print(f"  ✓ Образование: {len(education)} элементов")
    
    # Иностранный язык
    if general_info.get('foreign_language'):
        if fill_by_header(doc, ['иностранный язык', 'foreign language', 'foreign_language'], 
                         general_info['foreign_language'], 'foreign_language', debug=debug):
            replaced_count += 1
            print(f"  ✓ Иностранный язык: {general_info['foreign_language']}")
    
    # Гражданство / Локация
    if general_info.get('citizenship_location'):
        if fill_by_header(doc, ['гражданство', 'локация', 'citizenship', 'location', 'гражданство / локация'], 
                         general_info['citizenship_location'], 'citizenship_location', debug=debug):
            replaced_count += 1
            print(f"  ✓ Гражданство/Локация: {general_info['citizenship_location']}")
    
    # Занятость
    if general_info.get('employment'):
        if fill_by_header(doc, ['занятость', 'employment'], general_info['employment'], 'employment', debug=debug):
            replaced_count += 1
            print(f"  ✓ Занятость: {general_info['employment']}")
    
    # Статус
    if general_info.get('status'):
        if fill_by_header(doc, ['статус', 'status'], general_info['status'], 'status', debug=debug):
            replaced_count += 1
            print(f"  ✓ Статус: {general_info['status']}")
    
    # Скрининг
    screening = json_data.get('screening', {})
    
    # Hard skills
    hard_skills = screening.get('hard_skills', [])
    if hard_skills:
        if fill_list_by_header(doc, ['hard skills', 'hard_skills', 'hard'], hard_skills, 'hard_skills', debug=debug):
            replaced_count += 1
            print(f"  ✓ Hard skills: {len(hard_skills)} элементов")
    
    # Soft skills
    soft_skills = screening.get('soft_skills', [])
    if soft_skills:
        # Пробуем разные варианты поиска заголовка
        soft_keywords = [
            ['soft skills', 'soft_skills', 'soft'],  # Оригинальные
            ['soft skills:', 'soft_skills:', 'soft:'],  # С двоеточием
            ['soft skills 	', 'soft_skills 	'],  # С табуляцией (для таблиц)
        ]
        found = False
        for keywords in soft_keywords:
            if fill_list_by_header(doc, keywords, soft_skills, 'soft_skills', debug=debug):
                replaced_count += 1
                print(f"  ✓ Soft skills: {len(soft_skills)} элементов")
                found = True
                break
        if not found and debug:
            print(f"  ⚠️  Не найден заголовок для списка 'soft_skills'")
    
    # Опыт работы
    if work_experience:
        work_count = fill_work_experience_by_header(doc, work_experience)
        if work_count > 0:
            replaced_count += work_count
            print(f"  ✓ Опыт работы: {work_count} записей")
    
    # Проектный опыт - собираем из work_experience[].projects или создаем из work_experience
    all_projects = []
    
    # Собираем проекты из work_experience
    for work_item in work_experience:
        # Если есть проекты в work_experience, используем их
        projects = work_item.get('projects', [])
        if projects:
            for project in projects:
                # Преобразуем формат проекта из work_experience в формат project_experience
                project_data = {
                    'company': f"{work_item.get('company', '')} / {work_item.get('period', '')}",
                    'role': project.get('role', work_item.get('position', '')),
                    'tasks': project.get('tasks', []),
                    'technologies_and_tools': project.get('tools', project.get('technologies_and_tools', [])),
                    'achievements': project.get('achievements', [])
                }
                all_projects.append(project_data)
        else:
            # Если проектов нет, создаем проект из данных работы
            if work_item.get('company') or work_item.get('position'):
                project_data = {
                    'company': f"{work_item.get('company', '')} / {work_item.get('period', '')}",
                    'role': work_item.get('position', ''),
                    'tasks': work_item.get('responsibilities', []),
                    'technologies_and_tools': work_item.get('technologies', []),
                    'achievements': work_item.get('achievements', [])
                }
                all_projects.append(project_data)
    
    # Также добавляем проекты из project_experience (если есть)
    project_experience = json_data.get('project_experience', [])
    for project in project_experience:
        # Фильтруем плейсхолдеры
        company = project.get('company', '').strip()
        role = project.get('role', '').strip()
        if company not in ['Место работы / время', ''] and role not in ['Роль', '']:
            all_projects.append(project)
        elif project.get('tasks') and project.get('tasks') != ['Задачи']:
            all_projects.append(project)
        elif project.get('technologies_and_tools') and project.get('technologies_and_tools') != ['Технологии и инструменты']:
            all_projects.append(project)
    
    if all_projects:
        # Сортируем проекты по дате (от новых к старым)
        all_projects = sort_projects_by_date(all_projects)
        project_count = fill_project_experience_by_header(doc, all_projects)
        if project_count > 0:
            replaced_count += project_count
            print(f"  ✓ Проектный опыт: {project_count} записей")
    
    print(f"\nВсего заполнено полей: {replaced_count}")


def fill_list_in_table_column(doc, section_keywords, column_keywords, items, field_name, debug=False):
    """
    Заполняет список в таблице, где заголовки находятся в колонках.
    Например, в таблице "Скрининг" есть заголовок строки "СКРИНИНГ" и колонки "Hard skills" и "Soft skills".
    
    Args:
        doc: Документ
        section_keywords (list): Ключевые слова для поиска секции (например, ['скрининг', 'screening'])
        column_keywords (list): Ключевые слова для поиска колонки (например, ['hard skills', 'hard_skills'])
        items (list): Список элементов для заполнения
        field_name (str): Имя поля
        debug (bool): Выводить отладочную информацию
        
    Returns:
        bool: True если заполнение выполнено
    """
    # Ищем таблицу с секцией
    for table_idx, table in enumerate(doc.tables):
        section_row_idx = None
        
        # Ищем строку с заголовком секции
        for row_idx, row in enumerate(table.rows):
            row_text = ' '.join([p.text.strip() for cell in row.cells for p in cell.paragraphs]).lower()
            if any(kw.lower() in row_text for kw in section_keywords):
                section_row_idx = row_idx
                break
        
        if section_row_idx is None:
            continue
        
        # Ищем колонку с нужным заголовком в этой строке или в следующей строке
        target_cell = None
        header_row = table.rows[section_row_idx]
        
        # Сначала проверяем заголовки в той же строке
        for cell_idx, cell in enumerate(header_row.cells):
            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs]).lower()
            if any(kw.lower() in cell_text for kw in column_keywords):
                # Нашли заголовок колонки, берем следующую строку в этой колонке
                if section_row_idx + 1 < len(table.rows):
                    next_row = table.rows[section_row_idx + 1]
                    if cell_idx < len(next_row.cells):
                        target_cell = next_row.cells[cell_idx]
                        break
        
        # Если не нашли в той же строке, ищем в следующей строке (заголовки могут быть отдельной строкой)
        if target_cell is None and section_row_idx + 1 < len(table.rows):
            header_row_2 = table.rows[section_row_idx + 1]
            for cell_idx, cell in enumerate(header_row_2.cells):
                cell_text = ' '.join([p.text.strip() for p in cell.paragraphs]).lower()
                if any(kw.lower() in cell_text for kw in column_keywords):
                    # Нашли заголовок колонки, берем следующую строку в этой колонке
                    if section_row_idx + 2 < len(table.rows):
                        next_row = table.rows[section_row_idx + 2]
                        if cell_idx < len(next_row.cells):
                            target_cell = next_row.cells[cell_idx]
                            break
        
        if target_cell is None:
            continue
        
        # Заполняем ячейку списком элементов
        if debug:
            print(f"  🔍 Найдена ячейка таблицы для '{field_name}' в колонке {column_keywords[0]}")
        
        # Очищаем ячейку (удаляем все параграфы кроме первого)
        for para in target_cell.paragraphs[1:]:
            para._element.getparent().remove(para._element)
        
        # Заполняем список
        first_para = target_cell.paragraphs[0]
        first_para.clear()
        
        for i, item in enumerate(items):
            formatted_item = format_list_item(item)
            if i == 0:
                first_para.add_run(f"• {formatted_item}")
            else:
                new_para = target_cell.add_paragraph()
                new_para.add_run(f"• {formatted_item}")
        
        return True
    
    return False


def fill_list_by_header(doc, header_keywords, items, field_name, debug=False, use_bullets=True):
    """
    Заполняет список, ища его по заголовку.
    
    Args:
        doc: Документ
        header_keywords (list): Ключевые слова заголовка
        items (list): Список элементов
        field_name (str): Имя поля
        debug (bool): Выводить отладочную информацию
        
    Returns:
        bool: True если заполнение выполнено
    """
    # Для hard_skills и soft_skills сначала пробуем поиск в таблице Скрининг
    if field_name in ['hard_skills', 'soft_skills']:
        section_keywords = ['скрининг', 'screening']
        if fill_list_in_table_column(doc, section_keywords, header_keywords, items, field_name, debug=debug):
            return True
    
    # Обычный поиск по заголовкам
    target_info = find_empty_paragraph_after_header(doc, header_keywords)
    if target_info is None:
        if debug:
            print(f"  ⚠️  Не найден заголовок для списка '{field_name}' (ключевые слова: {header_keywords})")
        return False
    
    target_type, first_para = target_info
    
    # Находим индекс первого параграфа
    if target_type == 'paragraph':
        # Находим индекс параграфа в документе
        start_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para == first_para:
                start_idx = i
                break
        
        if start_idx is None:
            return False
        
        # Заполняем элементы списка
        for i, item in enumerate(items):
            formatted_item = format_list_item(item)
            insert_idx = start_idx + i
            if insert_idx >= len(doc.paragraphs):
                # Создаем новый параграф
                new_para = doc.paragraphs[-1].insert_paragraph_after()
                remove_paragraph_numbering(new_para)
                text_value = formatted_item if not use_bullets else f"• {formatted_item}"
                new_para.add_run(text_value)
            else:
                para = doc.paragraphs[insert_idx]
                old_text = para.text.strip()
                if not old_text or old_text in ['—', '-', '•', '']:
                    # Заполняем пустой параграф
                    para.clear()
                    remove_paragraph_numbering(para)
                    text_value = formatted_item if not use_bullets else f"• {formatted_item}"
                    para.add_run(text_value)
                else:
                    # Вставляем новый параграф перед текущим
                    new_para = para.insert_paragraph_before()
                    remove_paragraph_numbering(new_para)
                    text_value = formatted_item if not use_bullets else f"• {formatted_item}"
                    new_para.add_run(text_value)

        return True
    
    elif target_type == 'table_cell':
        # Заполняем в ячейке таблицы
        cell = first_para._parent  # Получаем ячейку из параграфа
        # Очищаем ячейку и добавляем элементы списка
        for para in cell.paragraphs[1:]:  # Удаляем все параграфы кроме первого
            para._element.getparent().remove(para._element)
        
        first_para.clear()
        remove_paragraph_numbering(first_para)
        for i, item in enumerate(items):
            formatted_item = format_list_item(item)
            if i == 0:
                text_value = formatted_item if not use_bullets else f"• {formatted_item}"
                first_para.add_run(text_value)
            else:
                new_para = cell.add_paragraph()
                remove_paragraph_numbering(new_para)
                text_value = formatted_item if not use_bullets else f"• {formatted_item}"
                new_para.add_run(text_value)

        return True

    return False


def fill_work_experience_by_header(doc, work_experience):
    """
    Заполняет опыт работы, ища секцию по заголовку.
    
    Args:
        doc: Документ
        work_experience (list): Список опыта работы
        
    Returns:
        int: Количество добавленных записей
    """
    header_info = find_section_by_header(doc, ['опыт работы', 'work experience', 'work_experience', 'опыт работы:'])
    if header_info is None:
        return 0
    
    header_type, header_idx = header_info
    
    # Работаем только с параграфами (не с таблицами для опыта работы)
    if header_type != 'paragraph':
        return 0
    
    # Если нет записей опыта работы, не заполняем
    if not work_experience:
        return 0
    
    # Находим место для вставки (после заголовка или после пустых строк)
    insert_idx = header_idx + 1
    while insert_idx < len(doc.paragraphs):
        para = doc.paragraphs[insert_idx]
        text = para.text.strip().lower()
        # Пропускаем пустые параграфы и заголовки других секций
        if not text or text in ['', '—', '-', '•']:
            insert_idx += 1
        elif any(kw in text for kw in ['проект', 'project', 'скрининг', 'screening', 'общая информация']):
                break
        else:
            insert_idx += 1
    
    # Если insert_idx вышел за границы, вставляем в конец документа
    if insert_idx >= len(doc.paragraphs):
        insert_idx = len(doc.paragraphs) - 1
        if insert_idx < 0:
            doc.add_paragraph()
            insert_idx = 0
    
    added_count = 0
    for work_item in work_experience:
        # Компания и позиция
        company = work_item.get('company', '')
        position = work_item.get('position', '')
        period = work_item.get('period', '')
        
        if company or position:
            # Проверяем границы
            if insert_idx >= len(doc.paragraphs):
                doc.add_paragraph()
                insert_idx = len(doc.paragraphs) - 1
            
            # Вставляем информацию о работе
            company_para = doc.paragraphs[insert_idx].insert_paragraph_before()
            if position:
                add_run_with_default_font(company_para, f"{position}")
                if company:
                    add_run_with_default_font(company_para, f" в {company}")
            else:
                add_run_with_default_font(company_para, company)
            insert_idx += 1
            
            # Период
            if period:
                if insert_idx >= len(doc.paragraphs):
                    doc.add_paragraph()
                    insert_idx = len(doc.paragraphs) - 1
                period_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                add_run_with_default_font(period_para, period)
                insert_idx += 1
            
            # Обязанности
            responsibilities = work_item.get('responsibilities', [])
            if responsibilities:
                for resp in responsibilities:
                    if insert_idx >= len(doc.paragraphs):
                        doc.add_paragraph()
                        insert_idx = len(doc.paragraphs) - 1
                    resp_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                    configure_bullet_paragraph(resp_para)
                    add_run_with_default_font(resp_para, f"• {resp}")
                    insert_idx += 1
            
            # Технологии
            technologies = work_item.get('technologies', [])
            if technologies:
                if insert_idx >= len(doc.paragraphs):
                    doc.add_paragraph()
                    insert_idx = len(doc.paragraphs) - 1
                tech_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                add_run_with_default_font(tech_para, f"Технологии: {', '.join(technologies)}")
                insert_idx += 1
            
            # Пустая строка между записями
            if added_count < len(work_experience) - 1:
                if insert_idx >= len(doc.paragraphs):
                    doc.add_paragraph()
                    insert_idx = len(doc.paragraphs) - 1
                doc.paragraphs[insert_idx].insert_paragraph_before()
                insert_idx += 1
        
        added_count += 1
    
    return added_count


def find_template_block_after_header(doc, header_idx, max_search=30):
    """
    Находит шаблонный блок после заголовка.
    Ищет блок, начинающийся с "Место работы / время" и содержащий поля "Роль:", "Задачи:", "Технологии и инструменты:".
    Включает пустые параграфы после меток, которые являются полями для заполнения.
    
    Args:
        doc: Документ
        header_idx (int): Индекс заголовка
        max_search (int): Максимальное количество параграфов для поиска
        
    Returns:
        list: Список параграфов шаблонного блока или None
    """
    template_block = []
    start_found = False
    found_fields = {'role': False, 'tasks': False, 'technologies': False}
    
    # Ищем начало блока - "Место работы / время"
    for i in range(header_idx + 1, min(header_idx + max_search, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # Проверяем, не заголовок ли это другой секции
        text_lower = text.lower()
        if any(kw in text_lower for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                           'work experience', 'general info', 'screening']):
            break
        
        # Ищем начало шаблона - "Место работы / время"
        if 'место работы / время' in text_lower or 'место работы' in text_lower:
            start_found = True
            template_block.append(para)
            # Продолжаем собирать блок до следующего "Место работы" или конца секции
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()
                next_text_lower = next_text.lower()
                
                # Если нашли следующее "Место работы" или заголовок другой секции, останавливаемся
                if ('место работы / время' in next_text_lower or 'место работы' in next_text_lower) and j > i + 1:
                    break
                if any(kw in next_text_lower for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                                  'work experience', 'general info', 'screening']):
                    break
                
                # Добавляем параграф в блок
                template_block.append(next_para)
                
                # Отслеживаем найденные поля
                if 'роль:' in next_text_lower:
                    found_fields['role'] = True
                if 'задачи:' in next_text_lower:
                    found_fields['tasks'] = True
                if 'технологии и инструменты' in next_text_lower or 'технологии:' in next_text_lower:
                    found_fields['technologies'] = True
                    # После технологий можем остановиться, если нашли все поля
                    if found_fields['role'] and found_fields['tasks'] and found_fields['technologies']:
                        # Добавляем еще один параграф после технологий (поле для значения)
                        if j + 1 < len(doc.paragraphs):
                            next_next_para = doc.paragraphs[j + 1]
                            next_next_text = next_next_para.text.strip().lower()
                            # Если следующий параграф не является меткой другого поля, добавляем его
                            if not any(kw in next_next_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                                template_block.append(next_next_para)
                    break
            break
    
    # Если нашли начало и хотя бы одно поле, возвращаем блок
    if start_found and len(template_block) >= 2:
        return template_block
    
    return None


def find_project_block_fields(doc, start_idx, max_search=20):
    """
    Находит поля в блоке проекта, начиная с указанного индекса.
    Возвращает словарь с индексами параграфов для каждого поля.
    
    Args:
        doc: Документ
        start_idx (int): Начальный индекс для поиска
        max_search (int): Максимальное количество параграфов для поиска
        
    Returns:
        dict: Словарь с ключами 'company', 'role_label', 'role_value', 'tasks_label', 
              'tasks_fields', 'tech_label', 'tech_value' и индексами параграфов
    """
    fields = {
        'company': None,
        'role_label': None,
        'role_value': None,
        'tasks_label': None,
        'tasks_fields': [],  # Список индексов для задач
        'achievements_label': None,
        'achievements_fields': [],
        'tech_label': None,
        'tech_value': None
    }
    
    for i in range(start_idx, min(start_idx + max_search, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        text_lower = text.lower()
        
        # Проверяем, не заголовок ли это другой секции
        if any(kw in text_lower for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                           'work experience', 'general info', 'screening']):
            break
        
        # Ищем "Место работы / время"
        if fields['company'] is None and ('место работы / время' in text_lower or 'место работы' in text_lower):
            fields['company'] = i
            continue
        
        # Ищем "Роль:"
        if fields['role_label'] is None and 'роль:' in text_lower:
            fields['role_label'] = i
            # Следующий параграф - это поле для значения роли
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip().lower()
                if not any(kw in next_text for kw in ['задачи:', 'технологии:', 'место работы', 'роль:']):
                    fields['role_value'] = i + 1
            continue
        
        # Ищем "Задачи:"
        if fields['tasks_label'] is None and 'задачи:' in text_lower:
            fields['tasks_label'] = i
            # Следующие параграфы до "Технологии" - это поля для задач
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip().lower()
                if 'достижения' in next_text or 'технологии' in next_text or 'место работы' in next_text:
                    break
                if next_text and not any(kw in next_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                    fields['tasks_fields'].append(j)
            continue

        # Ищем "Достижения:"
        if fields['achievements_label'] is None and 'достижения' in text_lower:
            fields['achievements_label'] = i
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip().lower()
                if 'технологии' in next_text or 'место работы' in next_text:
                    break
                if next_text and not any(kw in next_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы', 'достижения:']):
                    fields['achievements_fields'].append(j)
            continue
        
        # Ищем "Технологии и инструменты:"
        if fields['tech_label'] is None and ('технологии и инструменты' in text_lower or 'технологии:' in text_lower):
            fields['tech_label'] = i
            # Следующий параграф - это поле для значения технологий
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip().lower()
                if not any(kw in next_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                    fields['tech_value'] = i + 1
            break
    
    return fields


def find_all_project_blocks(doc, header_idx, max_search=200):
    """
    Находит все блоки проектов в секции, начиная с заголовка.
    Ищет блоки, начинающиеся с "Место работы / время" и содержащие поля:
    - Место работы / время
    - Роль:
    - Задачи:
    - Технологии и инструменты:
    
    Args:
        doc: Документ
        header_idx (int): Индекс заголовка секции
        max_search (int): Максимальное количество параграфов для поиска
        
    Returns:
        list: Список словарей с информацией о блоках (start_idx, end_idx, fields)
    """
    blocks = []
    # Начинаем поиск со следующего параграфа после заголовка
    current_idx = header_idx + 1
    
    while current_idx < len(doc.paragraphs) and current_idx < header_idx + max_search:
        para_text = doc.paragraphs[current_idx].text.strip()
        para_text_lower = para_text.lower()
        
        # Проверяем, не заголовок ли это другой секции
        if any(kw in para_text_lower for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                          'work experience', 'general info', 'screening']):
            break
    
        # Ищем начало блока - "Место работы / время"
        if 'место работы / время' in para_text_lower or 'место работы' in para_text_lower:
            print(f"     Найдено начало блока в параграфе {current_idx}: '{para_text[:50]}'")
            # Находим поля в этом блоке
            block_fields = find_project_block_fields(doc, current_idx)
            if block_fields['company'] is not None or block_fields['role_label'] is not None:
                # Определяем конец блока
                # Блок заканчивается перед следующим "Место работы / время" или перед заголовком другой секции
                block_end = current_idx
                found_technologies = False
                for i in range(current_idx, min(current_idx + 25, len(doc.paragraphs))):
                    next_para_text = doc.paragraphs[i].text.strip().lower()
                    
                    # Если нашли следующее "Место работы / время" - это начало следующего блока
                    if ('место работы / время' in next_para_text or 'место работы' in next_para_text) and i > current_idx:
                        block_end = i
                        break
                    
                    # Если нашли заголовок другой секции - останавливаемся
                    if any(kw in next_para_text for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                                          'work experience', 'general info', 'screening']):
                        block_end = i
                        break
                    
                    # Если нашли "Технологии" - это последнее поле блока
                    if 'технологии' in next_para_text and not found_technologies:
                        found_technologies = True
                        # Берем еще один параграф после технологий (поле для значения)
                        if i + 1 < len(doc.paragraphs):
                            # Проверяем, не является ли следующий параграф началом нового блока
                            next_next_text = doc.paragraphs[i + 1].text.strip().lower()
                            if 'место работы' not in next_next_text:
                                block_end = i + 2
                            else:
                                block_end = i + 1
                        else:
                            block_end = i + 1
                        # Продолжаем проверку, может быть еще параграфы в блоке
                        continue
                    
                    # Если уже нашли технологии, и следующий параграф не является меткой поля - это конец блока
                    if found_technologies:
                        if next_para_text and not any(kw in next_para_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                            # Это может быть пустая строка или текст, продолжаем
                            block_end = i + 1
                        elif not next_para_text:
                            # Пустой параграф - может быть разделитель между блоками
                            block_end = i + 1
                        else:
                            # Нашли метку нового поля - это начало следующего блока
                            if 'место работы' in next_para_text:
                                block_end = i
                                break
                    
                    block_end = i + 1
                
                blocks.append({
                    'start_idx': current_idx,
                    'end_idx': block_end,
                    'fields': block_fields
                })
                current_idx = block_end
                continue
        
        current_idx += 1
    
    return blocks


def find_all_project_blocks_in_tables(doc, header_idx):
    """
    Находит все блоки проектов в таблицах после заголовка.
    Ищет строки таблиц, содержащие "Место работы / время" и другие поля.
    
    Args:
        doc: Документ
        header_idx (int): Индекс заголовка секции
        
    Returns:
        list: Список словарей с информацией о блоках в таблицах
    """
    blocks = []
    
    # Ищем таблицы после заголовка
    # Находим, в какой таблице находится заголовок (если есть)
    header_table_idx = None
    header_row_idx = None
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_lower = para.text.lower().strip()
                    if 'проектный опыт' in text_lower or 'project experience' in text_lower:
                        header_table_idx = table_idx
                        header_row_idx = row_idx
                        break
                if header_table_idx is not None:
                    break
            if header_table_idx is not None:
                break
        if header_table_idx is not None:
            break
    
    # Если заголовок в таблице, ищем блоки в этой же таблице и следующих
    if header_table_idx is not None:
        print(f"     Заголовок найден в таблице {header_table_idx}, строке {header_row_idx}")
        # Ищем блоки в таблице, начиная со строки после заголовка
        table = doc.tables[header_table_idx]
        for row_idx in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[row_idx]
            # Проверяем первую ячейку строки на наличие "Место работы / время"
            if row.cells:
                first_cell_text = ' '.join([p.text.strip() for p in row.cells[0].paragraphs]).lower()
                if 'место работы / время' in first_cell_text or 'место работы' in first_cell_text:
                    # Нашли начало блока
                    block_info = {
                        'type': 'table',
                        'table_idx': header_table_idx,
                        'row_idx': row_idx,
                        'fields': find_project_block_fields_in_table_row(table, row_idx)
                    }
                    if block_info['fields']['company'] is not None or block_info['fields']['role_label'] is not None:
                        blocks.append(block_info)
                        print(f"     Найден блок в таблице {header_table_idx}, строке {row_idx}")
                        fields_info = block_info['fields']
                        print(f"        Поля: company={fields_info['company']}, role_label={fields_info['role_label']}, role_value={fields_info['role_value']}")
                        print(f"        tasks_label={fields_info['tasks_label']}, tasks_fields={len(fields_info['tasks_fields'])} полей")
                        print(f"        tech_label={fields_info['tech_label']}, tech_value={fields_info['tech_value']}")
                        
                        # Выводим детали найденных полей задач
                        if fields_info['tasks_fields']:
                            print(f"        Поля задач: {fields_info['tasks_fields'][:3]}...")  # Первые 3
    
    # Также ищем во всех таблицах (на случай, если структура другая)
    if not blocks:
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                if row.cells:
                    first_cell_text = ' '.join([p.text.strip() for p in row.cells[0].paragraphs]).lower()
                    if 'место работы / время' in first_cell_text or 'место работы' in first_cell_text:
                        block_info = {
                            'type': 'table',
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'fields': find_project_block_fields_in_table_row(table, row_idx)
                        }
                        if block_info['fields']['company'] is not None or block_info['fields']['role_label'] is not None:
                            blocks.append(block_info)
                            print(f"     Найден блок в таблице {table_idx}, строке {row_idx}")
    
    return blocks


def find_project_block_fields_in_table_row(table, start_row_idx):
    """
    Находит поля блока проекта в строке таблицы и следующих строках.
    Структура таблицы обычно такая:
    - Строка 1: "Место работы / время" | значение
    - Строка 2: "Роль:" | значение (или пустое)
    - Строка 3: "Задачи:" | задача 1
    - Строка 4: | задача 2 (или пустое)
    - Строка 5: "Технологии и инструменты:" | значение (или пустое)
    
    Args:
        table: Таблица
        start_row_idx (int): Индекс начальной строки блока
        
    Returns:
        dict: Словарь с полями блока
    """
    fields = {
        'company': None,
        'role_label': None,
        'role_value': None,
        'tasks_label': None,
        'tasks_fields': [],
        'achievements_label': None,
        'achievements_fields': [],
        'tech_label': None,
        'tech_value': None
    }
    
    # Отладочный вывод
    print(f"        🔍 Поиск полей в таблице, начиная со строки {start_row_idx}")
    print(f"           Всего строк в таблице: {len(table.rows)}")
    if len(table.rows) > 0:
        print(f"           Всего колонок в первой строке: {len(table.rows[0].cells)}")
    
    # Ищем в строке start_row_idx и следующих (максимум 15 строк для одного блока)
    for row_idx in range(start_row_idx, min(start_row_idx + 15, len(table.rows))):
        row = table.rows[row_idx]
        
        # Проверяем все ячейки в строке
        for cell_idx, cell in enumerate(row.cells):
            # Получаем весь текст ячейки
            cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()]).lower()
            if cell_text:
                print(f"           Строка {row_idx}, ячейка {cell_idx}: '{cell_text[:50]}'")
            
            # Место работы / время - обычно в первой строке блока
            if fields['company'] is None and ('место работы / время' in cell_text or 'место работы' in cell_text):
                fields['company'] = (row_idx, cell_idx)
                print(f"           ✓ Найдено 'Место работы' в ({row_idx}, {cell_idx})")
                # Значение может быть в следующей ячейке той же строки
                if cell_idx + 1 < len(row.cells):
                    next_cell = row.cells[cell_idx + 1]
                    next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                    # Если следующая ячейка не содержит метку другого поля, это значение
                    if next_cell_text and not any(kw in next_cell_text for kw in ['роль:', 'задачи:', 'технологии:', 'место работы']):
                        # Значение уже в ячейке, не нужно отдельно сохранять
                        pass
                continue
            
            # Роль: - обычно в отдельной строке после "Место работы"
            # ВАЖНО: может быть в одной ячейке с "Задачи:" и "Технологии:"
            if fields['role_label'] is None and 'роль:' in cell_text:
                fields['role_label'] = (row_idx, cell_idx)
                print(f"           ✓ Найдено 'Роль:' в ({row_idx}, {cell_idx})")
                
                # Если в ячейке также есть "Задачи:" и "Технологии:" - это особая структура
                # В этом случае значения находятся в следующих ячейках
                # Структура с 3 колонками: ячейка 0 = метки, ячейка 1 = значения, ячейка 2 = может быть пустая
                if 'задачи:' in cell_text and 'технологии' in cell_text:
                    print(f"              ⚠️ Все метки в одной ячейке! Ищу значения в следующих ячейках")
                    print(f"              Всего ячеек в строке: {len(row.cells)}")
                    
                    # В структуре с 3 колонками:
                    # - Ячейка 0: метки ("Роль: Задачи: Технологии:")
                    # - Ячейка 1: значения (роль, задачи, технологии - в разных параграфах или строках)
                    # - Ячейка 2: может быть пустая или с дополнительной информацией
                    
                    if cell_idx + 1 < len(row.cells):
                        # Значения находятся в следующей ячейке (ячейка 1)
                        next_cell = row.cells[cell_idx + 1]
                        next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                        print(f"              Проверяю ячейку ({row_idx}, {cell_idx + 1}): '{next_cell_text[:50] if next_cell_text else '(пустая)'}'")
                        print(f"              Количество параграфов в ячейке: {len(next_cell.paragraphs)}")
                        
                        # Показываем все параграфы
                        for para_idx, para in enumerate(next_cell.paragraphs):
                            para_text = para.text.strip()
                            print(f"                Параграф {para_idx}: '{para_text[:40] if para_text else '(пустой)'}'")
                        
                        # Роль обычно в первом параграфе ячейки 1
                        if len(next_cell.paragraphs) > 0:
                            first_para_text = next_cell.paragraphs[0].text.strip().lower()
                            if not first_para_text or first_para_text in ['', '—', '-', '•']:
                                # Пустой первый параграф - это поле для роли
                                fields['role_value'] = (row_idx, cell_idx + 1, 0)
                                print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1}, 0) - пустой параграф")
                            elif not any(kw in first_para_text for kw in ['задачи:', 'технологии:', 'место работы', 'роль:']):
                                # Первый параграф не содержит меток - это поле для роли
                                fields['role_value'] = (row_idx, cell_idx + 1, 0)
                                print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1}, 0)")
                            else:
                                # Первый параграф содержит метки, ищем дальше
                                fields['role_value'] = (row_idx, cell_idx + 1)
                                print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1})")
                        else:
                            # Нет параграфов - ячейка пустая, это поле для роли
                            fields['role_value'] = (row_idx, cell_idx + 1)
                            print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1}) - пустая ячейка")
                    else:
                        print(f"              ⚠️ Нет следующей ячейки! Всего ячеек в строке: {len(row.cells)}")
                    continue
                
                # Обычная структура - "Роль:" отдельно
                # Значение роли может быть:
                # 1. В следующей ячейке той же строки (если таблица в 2 колонки)
                if cell_idx + 1 < len(row.cells):
                    next_cell = row.cells[cell_idx + 1]
                    next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                    print(f"              Проверяю следующую ячейку ({row_idx}, {cell_idx + 1}): '{next_cell_text[:30]}'")
                    # Если ячейка пустая или содержит только пробелы/дефисы - это поле для значения
                    if not next_cell_text or next_cell_text in ['', '—', '-', '•']:
                        fields['role_value'] = (row_idx, cell_idx + 1)
                        print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1})")
                        continue
                    # Или если ячейка не содержит метку другого поля
                    elif not any(kw in next_cell_text for kw in ['задачи:', 'технологии:', 'место работы', 'роль:']):
                        fields['role_value'] = (row_idx, cell_idx + 1)
                        print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx + 1})")
                        continue
                
                # 2. В следующей строке той же колонки (если структура вертикальная)
                if row_idx + 1 < len(table.rows):
                    next_row = table.rows[row_idx + 1]
                    if cell_idx < len(next_row.cells):
                        next_row_cell = next_row.cells[cell_idx]
                        next_row_cell_text = ' '.join([p.text.strip() for p in next_row_cell.paragraphs if p.text.strip()]).lower()
                        print(f"              Проверяю следующую строку ({row_idx + 1}, {cell_idx}): '{next_row_cell_text[:30]}'")
                        # Если следующая строка пустая или не содержит метку другого поля - это поле для значения
                        if not next_row_cell_text or next_row_cell_text in ['', '—', '-', '•']:
                            fields['role_value'] = (row_idx + 1, cell_idx)
                            print(f"              ✓ Найдено поле для роли: ({row_idx + 1}, {cell_idx})")
                            continue
                        elif not any(kw in next_row_cell_text for kw in ['задачи:', 'технологии:', 'место работы', 'роль:']):
                            fields['role_value'] = (row_idx + 1, cell_idx)
                            print(f"              ✓ Найдено поле для роли: ({row_idx + 1}, {cell_idx})")
                            continue
                
                # 3. В следующем параграфе той же ячейки (если "Роль:" и значение в одной ячейке)
                if len(cell.paragraphs) > 1:
                    next_para_text = cell.paragraphs[1].text.strip().lower()
                    if next_para_text and not any(kw in next_para_text for kw in ['задачи:', 'технологии:', 'место работы', 'роль:']):
                        fields['role_value'] = (row_idx, cell_idx, 1)
                        print(f"              ✓ Найдено поле для роли: ({row_idx}, {cell_idx}, 1)")
                continue
            
            # Задачи: - обычно в отдельной строке после "Роль:"
            # ВАЖНО: задачи просто пишутся в ячейку (2, 1), не нужно искать отдельные поля
            if fields['tasks_label'] is None and 'задачи:' in cell_text:
                fields['tasks_label'] = (row_idx, cell_idx)
                print(f"           ✓ Найдено 'Задачи:' в ({row_idx}, {cell_idx})")
                
                # Если в ячейке также есть "Роль:" и "Технологии:" - это особая структура
                if 'роль:' in cell_text and 'технологии' in cell_text:
                    print(f"              ⚠️ Все метки в одной ячейке! Ищу значения задач в следующей ячейке")
                    # В структуре с 3 колонками задачи находятся в ячейке 1, начиная со второго параграфа
                    if cell_idx + 1 < len(row.cells):
                        next_cell = row.cells[cell_idx + 1]
                        next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                        print(f"              Проверяю ячейку ({row_idx}, {cell_idx + 1}) для задач: '{next_cell_text[:50] if next_cell_text else '(пустая)'}'")
                        print(f"              Количество параграфов в ячейке: {len(next_cell.paragraphs)}")
                        
                        # Проверяем параграфы в следующей ячейке
                        # Задачи обычно начиная со второго параграфа (индекс 1) или далее
                        # Первый параграф (индекс 0) обычно для роли
                        for para_idx, para in enumerate(next_cell.paragraphs):
                            para_text = para.text.strip().lower()
                            print(f"                Параграф {para_idx}: '{para_text[:30] if para_text else '(пустой)'}'")
                            # Пропускаем первый параграф (он для роли)
                            if para_idx > 0:
                                if para_text and not any(kw in para_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                                    fields['tasks_fields'].append((row_idx, cell_idx + 1, para_idx))
                                    print(f"              ✓ Найдено поле для задачи: ({row_idx}, {cell_idx + 1}, {para_idx})")
                        
                        # Также проверяем следующую строку (если есть)
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if cell_idx + 1 < len(next_row.cells):
                                next_row_cell = next_row.cells[cell_idx + 1]
                                next_row_cell_text = ' '.join([p.text.strip() for p in next_row_cell.paragraphs if p.text.strip()]).lower()
                                # Если следующая строка не содержит метки - это поле для задач
                                if not any(kw in next_row_cell_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                                    for para_idx, para in enumerate(next_row_cell.paragraphs):
                                        para_text = para.text.strip().lower()
                                        if para_text and not any(kw in para_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                                            fields['tasks_fields'].append((row_idx + 1, cell_idx + 1, para_idx))
                                            print(f"              ✓ Найдено поле для задачи: ({row_idx + 1}, {cell_idx + 1}, {para_idx})")
                                    # Если ячейка пустая, это тоже поле для задач
                                    if not next_row_cell_text:
                                        fields['tasks_fields'].append((row_idx + 1, cell_idx + 1))
                                        print(f"              ✓ Найдено пустое поле для задач: ({row_idx + 1}, {cell_idx + 1})")
                        
                        # Если ячейка пустая, добавляем её как поле для задач
                        if not next_cell_text or next_cell_text in ['', '—', '-', '•']:
                            fields['tasks_fields'].append((row_idx, cell_idx + 1))
                            print(f"              ✓ Найдено пустое поле для задач: ({row_idx}, {cell_idx + 1})")
                    else:
                        print(f"              ⚠️ Нет следующей ячейки! Всего ячеек в строке: {len(row.cells)}")
                    continue
                
                # Обычная структура: "Задачи:" в отдельной строке
                # Задачи просто пишутся в следующую ячейку (ячейка 1) той же строки
                if cell_idx + 1 < len(row.cells):
                    # Ячейка для задач - это просто ячейка (row_idx, cell_idx + 1)
                    fields['tasks_fields'].append((row_idx, cell_idx + 1))
                    print(f"              ✓ Найдено поле для задач: ({row_idx}, {cell_idx + 1}) - просто пишем в ячейку")
                else:
                    # Если нет следующей ячейки, ищем в следующей строке
                    if row_idx + 1 < len(table.rows):
                        next_row = table.rows[row_idx + 1]
                        if cell_idx < len(next_row.cells):
                            fields['tasks_fields'].append((row_idx + 1, cell_idx))
                            print(f"              ✓ Найдено поле для задач: ({row_idx + 1}, {cell_idx})")
                continue
            
            # Достижения
            if fields['achievements_label'] is None and 'достижения' in cell_text:
                fields['achievements_label'] = (row_idx, cell_idx)
                print(f"           ✓ Найдено 'Достижения:' в ({row_idx}, {cell_idx})")
                if cell_idx + 1 < len(row.cells):
                    fields['achievements_fields'].append((row_idx, cell_idx + 1))
                    print(f"              ✓ Поле достижений: ({row_idx}, {cell_idx + 1})")
                elif row_idx + 1 < len(table.rows):
                    next_row = table.rows[row_idx + 1]
                    if cell_idx < len(next_row.cells):
                        fields['achievements_fields'].append((row_idx + 1, cell_idx))
                        print(f"              ✓ Поле достижений: ({row_idx + 1}, {cell_idx})")
                continue

            # Технологии и инструменты: - обычно в последней строке блока
            # ВАЖНО: может быть в одной ячейке с "Роль:" и "Задачи:"
            if fields['tech_label'] is None and ('технологии и инструменты' in cell_text or 'технологии:' in cell_text):
                fields['tech_label'] = (row_idx, cell_idx)
                print(f"           ✓ Найдено 'Технологии:' в ({row_idx}, {cell_idx})")
                
                # Если в ячейке также есть "Роль:" и "Задачи:" - это особая структура
                if 'роль:' in cell_text and 'задачи:' in cell_text:
                    print(f"              ⚠️ Все метки в одной ячейке! Ищу значение технологий в следующей ячейке")
                    # В структуре с 3 колонками технологии находятся в ячейке 1, в последнем параграфе
                    # или в ячейке 2
                    if cell_idx + 1 < len(row.cells):
                        next_cell = row.cells[cell_idx + 1]
                        next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                        print(f"              Проверяю ячейку ({row_idx}, {cell_idx + 1}) для технологий: '{next_cell_text[:50] if next_cell_text else '(пустая)'}'")
                        print(f"              Количество параграфов в ячейке: {len(next_cell.paragraphs)}")
                        
                        # Проверяем последний параграф в ячейке (технологии обычно в конце)
                        if len(next_cell.paragraphs) > 1:
                            # Берем последний параграф для технологий (после роли и задач)
                            last_para_idx = len(next_cell.paragraphs) - 1
                            last_para_text = next_cell.paragraphs[last_para_idx].text.strip().lower()
                            print(f"                Последний параграф {last_para_idx}: '{last_para_text[:30] if last_para_text else '(пустой)'}'")
                            if not any(kw in last_para_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                                fields['tech_value'] = (row_idx, cell_idx + 1, last_para_idx)
                                print(f"              ✓ Найдено поле для технологий: ({row_idx}, {cell_idx + 1}, {last_para_idx})")
                            else:
                                # Или просто в ячейке
                                fields['tech_value'] = (row_idx, cell_idx + 1)
                                print(f"              ✓ Найдено поле для технологий: ({row_idx}, {cell_idx + 1})")
                        elif len(next_cell.paragraphs) == 1:
                            # Только один параграф - возможно, это только роль, технологии в другой ячейке
                            # Проверяем ячейку 2 (если есть)
                            if cell_idx + 2 < len(row.cells):
                                tech_cell = row.cells[cell_idx + 2]
                                tech_cell_text = ' '.join([p.text.strip() for p in tech_cell.paragraphs if p.text.strip()]).lower()
                                print(f"              Проверяю ячейку ({row_idx}, {cell_idx + 2}) для технологий: '{tech_cell_text[:50] if tech_cell_text else '(пустая)'}'")
                                if not tech_cell_text or not any(kw in tech_cell_text for kw in ['задачи:', 'роль:', 'технологии:', 'место работы']):
                                    fields['tech_value'] = (row_idx, cell_idx + 2)
                                    print(f"              ✓ Найдено поле для технологий: ({row_idx}, {cell_idx + 2})")
                        else:
                            # Нет параграфов - ячейка пустая
                            fields['tech_value'] = (row_idx, cell_idx + 1)
                            print(f"              ✓ Найдено поле для технологий: ({row_idx}, {cell_idx + 1}) - пустая ячейка")
                    else:
                        print(f"              ⚠️ Нет следующей ячейки! Всего ячеек в строке: {len(row.cells)}")
                    continue
                
                # Значение технологий может быть:
                # 1. В следующей ячейке той же строки
                if cell_idx + 1 < len(row.cells):
                    next_cell = row.cells[cell_idx + 1]
                    next_cell_text = ' '.join([p.text.strip() for p in next_cell.paragraphs if p.text.strip()]).lower()
                    # Если ячейка пустая или содержит только пробелы/дефисы - это поле для значения
                    if not next_cell_text or next_cell_text in ['', '—', '-', '•']:
                        fields['tech_value'] = (row_idx, cell_idx + 1)
                        break
                    # Или если ячейка не содержит метку другого поля
                    elif not any(kw in next_cell_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                        fields['tech_value'] = (row_idx, cell_idx + 1)
                        break
                
                # 2. В следующей строке той же колонки
                if row_idx + 1 < len(table.rows):
                    next_row = table.rows[row_idx + 1]
                    if cell_idx < len(next_row.cells):
                        next_row_cell = next_row.cells[cell_idx]
                        next_row_cell_text = ' '.join([p.text.strip() for p in next_row_cell.paragraphs if p.text.strip()]).lower()
                        # Если следующая строка пустая или не содержит метку другого поля - это поле для значения
                        if not next_row_cell_text or next_row_cell_text in ['', '—', '-', '•']:
                            fields['tech_value'] = (row_idx + 1, cell_idx)
                            break
                        elif not any(kw in next_row_cell_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                            fields['tech_value'] = (row_idx + 1, cell_idx)
                            break
                
                # 3. В следующем параграфе той же ячейки
                if len(cell.paragraphs) > 1:
                    next_para_text = cell.paragraphs[1].text.strip().lower()
                    if next_para_text and not any(kw in next_para_text for kw in ['место работы', 'роль:', 'задачи:', 'технологии:']):
                        fields['tech_value'] = (row_idx, cell_idx, 1)
                break
        
        # Если нашли технологии, останавливаемся (это последнее поле блока)
        if fields['tech_label'] is not None:
            break
    
    return fields


def fill_single_project_block(doc, block_fields, project_item):
    """
    Заполняет один блок проекта данными.
    
    Args:
        doc: Документ
        block_fields (dict): Словарь с полями блока
        project_item (dict): Данные проекта
        
    Returns:
        bool: True если блок заполнен успешно
    """
    company = project_item.get('company', '').strip()
    role = project_item.get('role', '').strip()
    tasks = project_item.get('tasks', [])
    achievements = project_item.get('achievements') or project_item.get('achievements_and_results', [])
    technologies = project_item.get('technologies_and_tools', [])
        
    # 1. Место работы / время
    if block_fields['company'] is not None:
        company_para = doc.paragraphs[block_fields['company']]
        if company and company != 'Место работы / время':
            replace_text_preserving_format(company_para, company_para.text, uppercase_duration_words(company))
    
    # 2. Роль
    if block_fields['role_label'] is not None:
        role_label_para = doc.paragraphs[block_fields['role_label']]
        # Убеждаемся, что метка "Роль:" есть
        if 'роль:' not in role_label_para.text.lower():
            replace_text_preserving_format(role_label_para, role_label_para.text, "Роль:", force_default_font=False)
    
    if block_fields['role_value'] is not None:
        role_value_para = doc.paragraphs[block_fields['role_value']]
        if role and role != 'Роль:':
            replace_text_preserving_format(role_value_para, role_value_para.text, role)
        else:
            replace_text_preserving_format(role_value_para, role_value_para.text, "")
        ensure_runs_not_bold(role_value_para)
    
    # 3. Задачи
    if block_fields['tasks_label'] is not None:
        tasks_label_para = doc.paragraphs[block_fields['tasks_label']]
        # Убеждаемся, что метка "Задачи:" есть
        if 'задачи:' not in tasks_label_para.text.lower():
            replace_text_preserving_format(tasks_label_para, tasks_label_para.text, "Задачи:", force_default_font=False)
    
    real_tasks = normalize_bullet_items(tasks, ['Задачи'])
    if real_tasks:
        if not set_bullet_list_in_document(doc, block_fields['tasks_fields'], real_tasks):
            pass

    real_achievements = normalize_bullet_items(achievements, ['Достижения'])
    if real_achievements:
        if not set_bullet_list_in_document(doc, block_fields['achievements_fields'], real_achievements):
            pass
    
    # 4. Технологии и инструменты
    if block_fields['tech_label'] is not None:
        tech_label_para = doc.paragraphs[block_fields['tech_label']]
        # Убеждаемся, что метка есть
        if 'технологии' not in tech_label_para.text.lower():
            replace_text_preserving_format(tech_label_para, tech_label_para.text, "Технологии и инструменты:", force_default_font=False)
    
    if block_fields['tech_value'] is not None:
        tech_value_para = doc.paragraphs[block_fields['tech_value']]
        if technologies and technologies != ['Технологии и инструменты']:
            real_tech = [t for t in technologies if t != 'Технологии и инструменты' and t.strip()]
            if real_tech:
                flat_tech = flatten_technology_entries(real_tech)
                tech_text = ', '.join(flat_tech) if flat_tech else ', '.join(real_tech)
                replace_text_preserving_format(tech_value_para, tech_value_para.text, tech_text)
            else:
                replace_text_preserving_format(tech_value_para, tech_value_para.text, "")
        else:
            replace_text_preserving_format(tech_value_para, tech_value_para.text, "")
        ensure_runs_not_bold(tech_value_para)
    
    return True


def fill_single_project_block_in_table(doc, block_info, project_item):
    """
    Заполняет один блок проекта в таблице данными.
    
    Args:
        doc: Документ
        block_info (dict): Информация о блоке в таблице
        project_item (dict): Данные проекта
        
    Returns:
        bool: True если блок заполнен успешно
    """
    table = doc.tables[block_info['table_idx']]
    fields = block_info['fields']
    
    company = project_item.get('company', '').strip()
    role = project_item.get('role', '').strip()
    tasks = project_item.get('tasks', [])
    technologies = project_item.get('technologies_and_tools', [])

    # Новый формат: таблица с одним столбцом (по строке на каждое поле)
    if all(len(row.cells) == 1 for row in table.rows):
        return fill_single_column_project_table(table, project_item)
    
    # Отладочный вывод
    print(f"  📝 Заполнение блока: {company}")
    print(f"     Роль: '{role}' (поле найдено: {fields['role_value'] is not None})")
    print(f"     Задачи: {len(tasks) if tasks else 0} (полей найдено: {len(fields['tasks_fields'])})")
    print(f"     Достижения: {len(achievements) if achievements else 0}")
    print(f"     Технологии: {len(technologies) if technologies else 0} (поле найдено: {fields['tech_value'] is not None})")
    
    # 1. Место работы / время
    if fields['company'] is not None:
        row_idx, cell_idx = fields['company']
        cell = table.rows[row_idx].cells[cell_idx]
        if company and company != 'Место работы / время':
            # Заменяем текст в ячейке
            if cell.paragraphs:
                # Заменяем текст в первом параграфе
                replace_text_preserving_format(
                    cell.paragraphs[0],
                    cell.paragraphs[0].text,
                    uppercase_duration_words(company)
                )
            else:
                cell.add_paragraph(uppercase_duration_words(company))
    
    # 2. Роль
    if fields['role_label'] is not None:
        row_idx, cell_idx = fields['role_label']
        cell = table.rows[row_idx].cells[cell_idx]
        # Убеждаемся, что метка "Роль:" есть
        cell_text = ' '.join([p.text.strip() for p in cell.paragraphs]).lower()
        if 'роль:' not in cell_text:
            if cell.paragraphs:
                replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, "Роль:", force_default_font=False)
    
    if fields['role_value'] is not None:
        # role_value может быть (row, cell) или (row, cell, para_index)
        if len(fields['role_value']) == 3:
            row_idx, cell_idx, para_idx = fields['role_value']
            cell = table.rows[row_idx].cells[cell_idx]
            if para_idx < len(cell.paragraphs):
                para = cell.paragraphs[para_idx]
                if role and role != 'Роль:':
                    print(f"     ✓ Заполняю роль в ячейке ({row_idx}, {cell_idx}), параграф {para_idx}")
                    replace_text_preserving_format(para, para.text, role)
                else:
                    replace_text_preserving_format(para, para.text, "")
                ensure_runs_not_bold(para)
            else:
                if role and role != 'Роль:':
                    print(f"     ✓ Добавляю роль в ячейку ({row_idx}, {cell_idx})")
                    new_para = cell.add_paragraph("")
                    add_run_with_default_font(new_para, role)
        else:
            row_idx, cell_idx = fields['role_value']
            cell = table.rows[row_idx].cells[cell_idx]
            if role and role != 'Роль:':
                print(f"     ✓ Заполняю роль в ячейке ({row_idx}, {cell_idx})")
                if cell.paragraphs:
                    # Если ячейка пустая или содержит только пробелы, заменяем
                    current_text = cell.paragraphs[0].text.strip()
                    replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, role)
                    ensure_runs_not_bold(cell.paragraphs[0])
                else:
                    new_para = cell.add_paragraph("")
                    add_run_with_default_font(new_para, role)
            else:
                if cell.paragraphs:
                    replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, "")
                    ensure_runs_not_bold(cell.paragraphs[0])
    else:
        print(f"     ⚠️ Поле для роли не найдено!")
    
    # 3. Задачи
    if fields['tasks_label'] is not None:
        row_idx, cell_idx = fields['tasks_label']
        label_cell = table.rows[row_idx].cells[cell_idx]
        cell_text = ' '.join([p.text.strip() for p in label_cell.paragraphs]).lower()
        if 'задачи:' not in cell_text:
            if label_cell.paragraphs:
                replace_text_preserving_format(label_cell.paragraphs[0], label_cell.paragraphs[0].text, "Задачи:", force_default_font=False)
    real_tasks = normalize_bullet_items(tasks, ['Задачи'])
    if real_tasks:
        task_cell = None
        if fields['tasks_fields']:
            task_field = fields['tasks_fields'][0]
            if len(task_field) == 3:
                row_idx, cell_idx, _ = task_field
            else:
                row_idx, cell_idx = task_field
            task_cell = table.rows[row_idx].cells[cell_idx]
        elif fields['tasks_label'] is not None:
            row_idx, cell_idx = fields['tasks_label']
            if cell_idx + 1 < len(table.rows[row_idx].cells):
                task_cell = table.rows[row_idx].cells[cell_idx + 1]
        if task_cell is not None:
            set_bullet_list_in_cell(task_cell, real_tasks)
        else:
            print(f"     ⚠️ Не найдено поле для задач!")

    if fields.get('achievements_label') is not None:
        row_idx, cell_idx = fields['achievements_label']
        label_cell = table.rows[row_idx].cells[cell_idx]
        cell_text = ' '.join([p.text.strip() for p in label_cell.paragraphs]).lower()
        if 'достижения' not in cell_text:
            if label_cell.paragraphs:
                replace_text_preserving_format(label_cell.paragraphs[0], label_cell.paragraphs[0].text, "Достижения:", force_default_font=False)
    real_achievements = normalize_bullet_items(achievements, ['Достижения'])
    if real_achievements:
        ach_cell = None
        if fields.get('achievements_fields'):
            ach_field = fields['achievements_fields'][0]
            if len(ach_field) == 3:
                row_idx, cell_idx, _ = ach_field
            else:
                row_idx, cell_idx = ach_field
            ach_cell = table.rows[row_idx].cells[cell_idx]
        elif fields.get('achievements_label') is not None:
            row_idx, cell_idx = fields['achievements_label']
            if cell_idx + 1 < len(table.rows[row_idx].cells):
                ach_cell = table.rows[row_idx].cells[cell_idx + 1]
        if ach_cell is not None:
            set_bullet_list_in_cell(ach_cell, real_achievements)
    
    # 4. Технологии и инструменты
    if fields['tech_label'] is not None:
        row_idx, cell_idx = fields['tech_label']
        cell = table.rows[row_idx].cells[cell_idx]
        # Убеждаемся, что метка есть
        cell_text = ' '.join([p.text.strip() for p in cell.paragraphs]).lower()
        if 'технологии' not in cell_text:
            if cell.paragraphs:
                replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, "Технологии и инструменты:")
    
    if fields['tech_value'] is not None:
        if len(fields['tech_value']) == 3:
            row_idx, cell_idx, para_idx = fields['tech_value']
            cell = table.rows[row_idx].cells[cell_idx]
            if para_idx < len(cell.paragraphs):
                para = cell.paragraphs[para_idx]
                if technologies and technologies != ['Технологии и инструменты']:
                    real_tech = [t for t in technologies if t != 'Технологии и инструменты' and t.strip()]
                    if real_tech:
                        flat_tech = flatten_technology_entries(real_tech)
                        tech_text = ', '.join(flat_tech) if flat_tech else ', '.join(real_tech)
                        print(f"     ✓ Заполняю технологии в ячейке ({row_idx}, {cell_idx}), параграф {para_idx}")
                        replace_text_preserving_format(para, para.text, tech_text)
                        ensure_runs_not_bold(para)
                    else:
                        replace_text_preserving_format(para, para.text, "")
                else:
                    replace_text_preserving_format(para, para.text, "")
            else:
                if technologies and technologies != ['Технологии и инструменты']:
                    real_tech = [t for t in technologies if t != 'Технологии и инструменты' and t.strip()]
                    if real_tech:
                        flat_tech = flatten_technology_entries(real_tech)
                        tech_text = ', '.join(flat_tech) if flat_tech else ', '.join(real_tech)
                        print(f"     ✓ Добавляю технологии в ячейку ({row_idx}, {cell_idx})")
                        new_para = cell.add_paragraph("")
                        add_run_with_default_font(new_para, tech_text)
        else:
            row_idx, cell_idx = fields['tech_value']
            cell = table.rows[row_idx].cells[cell_idx]
            if technologies and technologies != ['Технологии и инструменты']:
                real_tech = [t for t in technologies if t != 'Технологии и инструменты' and t.strip()]
                if real_tech:
                    flat_tech = flatten_technology_entries(real_tech)
                    tech_text = ', '.join(flat_tech) if flat_tech else ', '.join(real_tech)
                    print(f"     ✓ Заполняю технологии в ячейке ({row_idx}, {cell_idx})")
                    if cell.paragraphs:
                        replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, tech_text)
                        ensure_runs_not_bold(cell.paragraphs[0])
                    else:
                        new_para = cell.add_paragraph("")
                        add_run_with_default_font(new_para, tech_text)
                else:
                    if cell.paragraphs:
                        replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, "")
            else:
                if cell.paragraphs:
                    replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, "")
    else:
        print(f"     ⚠️ Поле для технологий не найдено!")
    
    return True


def fill_single_column_project_table(table, project_item):
    """Заполняет блок проекта в таблице с одним столбцом."""
    def find_row(keyword):
        keyword = keyword.lower()
        for idx, row in enumerate(table.rows):
            if not row.cells:
                continue
            cell_text = ' '.join(p.text.strip().lower() for p in row.cells[0].paragraphs)
            if keyword in cell_text:
                return idx
        return None

    def get_cell(row_idx):
        if row_idx is None or row_idx >= len(table.rows):
            return None
        return table.rows[row_idx].cells[0] if table.rows[row_idx].cells else None

    def set_label_value(cell, fallback_label, value):
        if cell is None:
            return False
        if not cell.paragraphs:
            cell.add_paragraph('')
        first_para = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
        text = first_para.text.strip()
        colon_idx = text.find(':')
        label_text = fallback_label
        if colon_idx != -1:
            label_text = text[:colon_idx + 1]
        write_label_and_value(first_para, label_text, value.strip() if value else '')
        return True

    # Место работы / время
    company_value = project_item.get('company', '').strip()
    if company_value and company_value != 'Место работы / время':
        company_value = uppercase_duration_words(company_value)
        cell = get_cell(find_row('место работы'))
        if cell and cell.paragraphs:
            replace_text_preserving_format(cell.paragraphs[0], cell.paragraphs[0].text, company_value)

    # Роль
    role_value = project_item.get('role', '').strip()
    set_label_value(get_cell(find_row('роль')), 'Роль:', role_value)

    # Задачи
    tasks_items = normalize_bullet_items(project_item.get('tasks', []), ['Задачи'])
    set_labeled_bullet_list(get_cell(find_row('задачи')), 'Задачи:', tasks_items)

    # Достижения
    achievements_items = normalize_bullet_items(
        project_item.get('achievements') or project_item.get('achievements_and_results', []),
        ['Достижения']
    )
    set_labeled_bullet_list(get_cell(find_row('достижения')), 'Достижения:', achievements_items)

    # Технологии и инструменты
    tech_items = normalize_bullet_items(project_item.get('technologies_and_tools', []), ['Технологии и инструменты'])
    tech_text = '; '.join(tech_items)
    set_label_value(get_cell(find_row('технологии')), 'Технологии и инструменты:', tech_text)

    return True


def parse_date_from_period(period_str):
    """
    Парсит дату из строки периода.
    Форматы: "Январь 2025 — настоящее время", "Январь 2019 — настоящее время", 
             "Январь 2018 — Октябрь 2018", "Май 2014 — Январь 2017"
    
    Args:
        period_str (str): Строка с периодом
        
    Returns:
        tuple: (год, месяц) для сортировки, или (0, 0) если не удалось распарсить
    """
    if not period_str:
        return (0, 0)
    
    # Словарь месяцев
    months = {
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4, 'май': 5, 'июнь': 6,
        'июль': 7, 'август': 8, 'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    
    period_lower = period_str.lower()
    
    # Ищем первую дату (начало периода)
    # Паттерн: "месяц год" или "месяц YYYY"
    match = re.search(r'(\w+)\s+(\d{4})', period_lower)
    if match:
        month_name = match.group(1)
        year = int(match.group(2))
        month = months.get(month_name, 0)
        if month > 0:
            return (year, month)
    
    # Если не нашли, пытаемся найти только год
    year_match = re.search(r'(\d{4})', period_lower)
    if year_match:
        year = int(year_match.group(1))
        return (year, 0)
    
    return (0, 0)


def sort_projects_by_date(projects):
    """
    Сортирует проекты по дате начала (от новых к старым).
    
    Args:
        projects (list): Список проектов
        
    Returns:
        list: Отсортированный список проектов
    """
    def get_sort_key(project):
        # Извлекаем период из поля company
        company = project.get('company', '')
        period = ''
        
        # Период может быть в разных форматах:
        # 1. "Компания / Период"
        if ' / ' in company:
            period = company.split(' / ', 1)[1]
        # 2. "Компания, Период" - период после запятой
        elif ', ' in company:
            # Ищем последнюю запятую с датой после неё
            parts = company.split(', ')
            if len(parts) > 1:
                # Берем последнюю часть после запятой
                potential_period = parts[-1]
                # Проверяем, содержит ли это дату (год)
                if re.search(r'\d{4}', potential_period):
                    period = potential_period
        # 3. "Компания (Период)" - период в скобках (может быть несколько скобок)
        elif '(' in company and ')' in company:
            # Ищем последние скобки с периодом (обычно это последние скобки в строке)
            # Ищем все скобки
            matches = list(re.finditer(r'\(([^)]+)\)', company))
            if matches:
                # Берем последние скобки (обычно там период)
                last_match = matches[-1]
                period = last_match.group(1)
                # Проверяем, содержит ли это дату (год)
                if not re.search(r'\d{4}', period):
                    # Если нет года, пробуем предпоследние скобки
                    if len(matches) > 1:
                        period = matches[-2].group(1)
        # 4. Отдельное поле period
        if not period:
            period = project.get('period', '')
        
        # Парсим дату
        year, month = parse_date_from_period(period)
        # Отладочный вывод для каждого проекта
        print(f"        Проект: '{company[:50]}' -> Период: '{period[:40]}' -> Дата: ({year}, {month})")
        # Возвращаем кортеж для сортировки (отрицательные значения для сортировки по убыванию)
        # Используем большие числа для проектов без даты, чтобы они были в конце
        return (-year if year > 0 else 9999, -month if month > 0 else 0)
    
    # Сортируем проекты
    sorted_projects = sorted(projects, key=get_sort_key)
    
    # Отладочный вывод
    print(f"  📅 Сортировка проектов по дате (от новых к старым):")
    for i, project in enumerate(sorted_projects):
        company = project.get('company', '')
        print(f"     {i+1}. {company[:60]}")
    
    return sorted_projects


def fill_project_experience_by_header(doc, project_experience):
    """
    Заполняет проектный опыт, находя все блоки-таблички в секции и заполняя их данными.
    Если блоков больше чем проектов - оставляет лишние пустыми.
    Если проектов больше чем блоков - создает новые блоки.
    
    Структура: ПРОЕКТНЫЙ ОПЫТ (заголовок)
              Место работы / время (блок 1)
              Роль:
              Задачи:
              Технологии и инструменты:
              Место работы / время (блок 2)
              ...
    
    Args:
        doc: Документ
        project_experience (list): Список проектного опыта
        
    Returns:
        int: Количество заполненных записей
    """
    # Ищем заголовок "ПРОЕКТНЫЙ ОПЫТ" (в разных вариантах написания)
    header_keywords = ['проектный опыт', 'project experience', 'project_experience', 'проектный опыт:']
    
    # Отладка: выводим первые параграфы документа
    print("\n🔍 Поиск секции 'Проектный опыт'...")
    print(f"   Всего параграфов в документе: {len(doc.paragraphs)}")
    print("   Первые 30 параграфов:")
    for i in range(min(30, len(doc.paragraphs))):
        para_text = doc.paragraphs[i].text.strip()
        if para_text:
            print(f"   [{i:2d}] {para_text[:60]}")
    
    header_info = find_section_by_header(doc, header_keywords)
    if header_info is None:
        print("  ⚠️  Не найден заголовок 'Проектный опыт'")
        print(f"     Искали ключевые слова: {header_keywords}")
        return 0
    
    header_type, header_idx = header_info
    print(f"  ✓ Найден заголовок: тип={header_type}, индекс={header_idx}")
    
    # Работаем только с параграфами (не с таблицами для проектного опыта)
    if header_type != 'paragraph':
        print(f"  ⚠️  Заголовок найден в таблице, а не в параграфе")
        return 0
    
    # header_idx - это индекс следующего параграфа после заголовка
    # Нам нужен индекс самого заголовка, поэтому уменьшаем на 1
    actual_header_idx = header_idx - 1 if header_idx > 0 else 0
    
    if actual_header_idx < len(doc.paragraphs):
        header_text = doc.paragraphs[actual_header_idx].text.strip()
        print(f"  ✓ Текст заголовка: '{header_text}'")
    
    # Фильтруем записи проекта, которые не являются плейсхолдерами
    real_projects = []
    for p in project_experience:
        company = p.get('company', '').strip()
        role = p.get('role', '').strip()
        # Пропускаем плейсхолдеры
        if company not in ['Место работы / время', ''] and role not in ['Роль', '']:
            real_projects.append(p)
        # Также проверяем, есть ли реальные задачи или технологии
        elif p.get('tasks') and p.get('tasks') != ['Задачи']:
            real_projects.append(p)
        elif p.get('technologies_and_tools') and p.get('technologies_and_tools') != ['Технологии и инструменты']:
            real_projects.append(p)
    
    if not real_projects:
        return 0
    
    # Находим все существующие блоки в секции (начинаем поиск после заголовка)
    print(f"\n🔍 Поиск блоков проектов после параграфа {actual_header_idx}...")
    print(f"   Параграфы после заголовка (первые 20):")
    for i in range(actual_header_idx + 1, min(actual_header_idx + 21, len(doc.paragraphs))):
        para_text = doc.paragraphs[i].text.strip()
        if para_text:
            print(f"   [{i:2d}] {para_text[:60]}")
    
    # Сначала ищем в параграфах
    existing_blocks = find_all_project_blocks(doc, actual_header_idx)
    
    # Если не нашли в параграфах, ищем в таблицах
    if not existing_blocks:
        print("   Блоки не найдены в параграфах, ищем в таблицах...")
        print(f"   Всего таблиц в документе: {len(doc.tables)}")
        existing_blocks = find_all_project_blocks_in_tables(doc, actual_header_idx)
    
    if not existing_blocks:
        print("  ⚠️  Не найдено блоков проектов в секции")
        print(f"     Искали блоки после параграфа {actual_header_idx}")
        print(f"     Проверьте, что в документе есть блоки, начинающиеся с 'Место работы / время'")
        return 0
    
    print(f"  ✓ Найдено блоков проектов: {len(existing_blocks)}")
    for i, block in enumerate(existing_blocks):
        if block.get('type') == 'table':
            print(f"     Блок {i+1}: таблица {block['table_idx']}, строка {block['row_idx']}")
        else:
            print(f"     Блок {i+1}: параграфы {block.get('start_idx', '?')}-{block.get('end_idx', '?')}")
    
    # Сохраняем информацию о первом блоке для клонирования (если нужно)
    first_block = existing_blocks[0]
    template_block_length = None
    template_para_texts = []
    
    if first_block.get('type') != 'table':
        # Для параграфов сохраняем тексты
        template_block_length = first_block['end_idx'] - first_block['start_idx']
        for i in range(template_block_length):
            if first_block['start_idx'] + i < len(doc.paragraphs):
                template_para_texts.append(doc.paragraphs[first_block['start_idx'] + i].text)
            else:
                template_para_texts.append("")
    
    # Заполняем существующие блоки
    filled_count = 0
    for block_idx, block in enumerate(existing_blocks):
        if block_idx < len(real_projects):
            # Заполняем блок данными проекта
            project_item = real_projects[block_idx]
            if block.get('type') == 'table':
                fill_single_project_block_in_table(doc, block, project_item)
            else:
                fill_single_project_block(doc, block['fields'], project_item)
            filled_count += 1
            company = project_item.get('company', 'Не указано')
            role = project_item.get('role', 'Не указано')
            print(f"  ✓ Проект {block_idx + 1}: {company} - {role}")
    
    # Если проектов больше чем блоков, создаем новые блоки
    if len(real_projects) > len(existing_blocks):
        print(f"  ⚠️  Проектов ({len(real_projects)}) больше чем блоков ({len(existing_blocks)})")
        print(f"     Создание новых блоков пока не поддерживается для таблиц")
        # TODO: Реализовать клонирование блоков в таблицах
    
    return filled_count


def fill_project_experience_simple(doc, header_idx, real_projects):
    """
    Простой метод заполнения проектного опыта (без клонирования шаблона).
    Используется, если шаблонный блок не найден.
    """
    # Находим место для вставки
    insert_idx = header_idx + 1
    while insert_idx < len(doc.paragraphs):
        para = doc.paragraphs[insert_idx]
        text = para.text.strip().lower()
        if not text or text in ['', '—', '-', '•']:
            insert_idx += 1
        else:
            if any(kw in text for kw in ['опыт работы', 'общая информация', 'скрининг', 
                                         'work experience', 'general info', 'screening']):
                break
            insert_idx += 1
    
    if insert_idx >= len(doc.paragraphs):
        insert_idx = len(doc.paragraphs) - 1
        if insert_idx < 0:
            doc.add_paragraph()
            insert_idx = 0
    
    added_count = 0
    for project_item in real_projects:
        company = project_item.get('company', '').strip()
        role = project_item.get('role', '').strip()
        tasks = project_item.get('tasks', [])
        technologies = project_item.get('technologies_and_tools', [])
        
        if not company and not role and not tasks and not technologies:
            continue
        
        if insert_idx >= len(doc.paragraphs):
            doc.add_paragraph()
            insert_idx = len(doc.paragraphs) - 1
        
        if company and company != 'Место работы / время':
            company_para = doc.paragraphs[insert_idx].insert_paragraph_before()
            add_run_with_default_font(company_para, uppercase_duration_words(company))
            insert_idx += 1
        
        if insert_idx >= len(doc.paragraphs):
            doc.add_paragraph()
            insert_idx = len(doc.paragraphs) - 1
        
        role_para = doc.paragraphs[insert_idx].insert_paragraph_before()
        label_run = add_run_with_default_font(role_para, "Роль:")
        label_run.bold = True
        if role and role != 'Роль:':
            add_run_with_default_font(role_para, f" {role}")
        insert_idx += 1
        
        if tasks and tasks != ['Задачи']:
            real_tasks = [t for t in tasks if t != 'Задачи' and t.strip()]
            if real_tasks:
                for task in real_tasks:
                    if insert_idx >= len(doc.paragraphs):
                        doc.add_paragraph()
                        insert_idx = len(doc.paragraphs) - 1
                    task_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                    configure_bullet_paragraph(task_para)
                    add_run_with_default_font(task_para, f"• {task}")
                    insert_idx += 1
        
        if insert_idx >= len(doc.paragraphs):
            doc.add_paragraph()
            insert_idx = len(doc.paragraphs) - 1
        
        if technologies and technologies != ['Технологии и инструменты']:
            real_tech = [t for t in technologies if t != 'Технологии и инструменты' and t.strip()]
            if real_tech:
                tech_para = doc.paragraphs[insert_idx].insert_paragraph_before()
                tech_para.clear()
                if tech_para._element.pPr is not None:
                    numPr = tech_para._element.pPr.numPr
                    if numPr is not None:
                        tech_para._element.pPr.remove(numPr)
                # Если технологии уже в формате "Категория: технологии", используем как есть
                # Иначе добавляем общую категорию
                if any(':' in t for t in real_tech):
                    tech_text = '\n'.join(real_tech)
                    add_run_with_default_font(tech_para, tech_text)
                else:
                    tech_text = ', '.join(real_tech)
                    add_run_with_default_font(tech_para, f"Технологии и инструменты: {tech_text}")
                insert_idx += 1
        
        if added_count < len(real_projects) - 1:
            if insert_idx >= len(doc.paragraphs):
                doc.add_paragraph()
                insert_idx = len(doc.paragraphs) - 1
            doc.paragraphs[insert_idx].insert_paragraph_before()
            insert_idx += 1
        
        added_count += 1
    
    return added_count


def main():
    """Основная функция."""
    parser = argparse.ArgumentParser(
        description="Заполнение Word документа данными из JSON",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  python json_to_docx.py data.json
  python json_to_docx.py data.json --template example_cv_docx.docx
  python json_to_docx.py data.json --output result.docx
  python json_to_docx.py data.json --template template.docx --output result.docx

Плейсхолдеры в шаблоне:
  Простые поля: {{vacancy}}, {{pitch}}, {{foreign_language}}, и т.д.
  Списки: {{#skills_and_tools}}...{{/skills_and_tools}}
  Блоки: {{#work_experience}}...{{/work_experience}}
        """
    )
    
    parser.add_argument("json_file", help="Путь к JSON файлу с данными")
    parser.add_argument(
        "--template", "-t",
        default="parser/template/example_cv_docx.docx",
        help="Путь к шаблону Word (по умолчанию: parser/template/example_cv_docx.docx)"
    )
    parser.add_argument(
        "--output", "-o",
        help="Путь к выходному файлу (по умолчанию: <имя_json>_filled.docx)"
    )
    
    args = parser.parse_args()
    
    # Проверка JSON файла
    if not os.path.exists(args.json_file):
        print(f"Ошибка: файл '{args.json_file}' не найден.")
        sys.exit(1)
    
    # Проверка шаблона
    if not os.path.exists(args.template):
        print(f"Ошибка: шаблон '{args.template}' не найден.")
        sys.exit(1)
    
    # Определение выходного файла
    if args.output:
        output_path = args.output
    else:
        json_file = Path(args.json_file)
        output_path = json_file.stem + "_filled.docx"
    
    # Загрузка JSON
    print(f"Загрузка JSON: {args.json_file}")
    json_data = load_json(args.json_file)
    
    # Заполнение документа
    fill_document(args.template, json_data, output_path)


if __name__ == "__main__":
    main()
