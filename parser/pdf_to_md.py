#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для извлечения текста из PDF файла и сохранения в Markdown файл.
"""

import sys
import os
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    print("Ошибка: библиотека pdfplumber не установлена.")
    print("Установите её командой: pip install pdfplumber")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    # DOCX поддержка опциональна, если не установлена - просто не будет работать
    Document = None



def extract_text_from_pdf(pdf_path):
    """
    Извлекает весь текст из PDF файла.
    
    Args:
        pdf_path (str): Путь к PDF файлу
        
    Returns:
        str: Извлеченный текст
    """
    text_content = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            print(f"Обработка PDF файла: {pdf_path}")
            print(f"Количество страниц: {len(pdf.pages)}")
            
            for i, page in enumerate(pdf.pages, 1):
                print(f"Обработка страницы {i}/{len(pdf.pages)}...")
                text = page.extract_text()
                if text:
                    text_content.append(text)
                    
    except Exception as e:
        print(f"Ошибка при чтении PDF файла: {e}")
        sys.exit(1)
    
    return "\n\n".join(text_content)


def extract_text_from_docx(docx_path):
    """
    Извлекает весь текст из DOCX файла.
    
    Args:
        docx_path (str): Путь к DOCX файлу
        
    Returns:
        str: Извлеченный текст
    """
    if Document is None:
        print("Ошибка: библиотека python-docx не установлена.")
        print("Установите её командой: pip install python-docx")
        sys.exit(1)
    
    text_content = []
    
    try:
        doc = Document(docx_path)
        print(f"Обработка DOCX файла: {docx_path}")
        print(f"Количество параграфов: {len(doc.paragraphs)}")
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_content.append(text)
                
    except Exception as e:
        print(f"Ошибка при чтении DOCX файла: {e}")
        sys.exit(1)
    
    return "\n\n".join(text_content)


def save_to_markdown(text, output_path):
    """
    Сохраняет текст в Markdown файл.
    
    Args:
        text (str): Текст для сохранения
        output_path (str): Путь к выходному MD файлу
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"\nТекст успешно сохранен в: {output_path}")
    except Exception as e:
        print(f"Ошибка при сохранении файла: {e}")
        sys.exit(1)


def main():
    """Основная функция."""
    if len(sys.argv) < 2:
        print("Использование: python pdf_to_md.py <путь_к_pdf_файлу> [путь_к_выходному_md_файлу]")
        print("\nПримеры:")
        print("  python pdf_to_md.py document.pdf")
        print("  python pdf_to_md.py document.pdf output.md")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    
    # Проверка существования PDF файла
    if not os.path.exists(pdf_path):
        print(f"Ошибка: файл '{pdf_path}' не найден.")
        sys.exit(1)
    
    # Определение пути к выходному файлу
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # Если выходной файл не указан, создаем его на основе имени PDF
        pdf_file = Path(pdf_path)
        output_path = pdf_file.with_suffix('.md')
    
    # Извлечение текста
    text = extract_text_from_pdf(pdf_path)
    
    if not text.strip():
        print("Предупреждение: не удалось извлечь текст из PDF файла.")
        print("Возможно, PDF файл содержит только изображения или защищен от копирования.")
    
    # Сохранение в Markdown
    save_to_markdown(text, output_path)


if __name__ == "__main__":
    main()

